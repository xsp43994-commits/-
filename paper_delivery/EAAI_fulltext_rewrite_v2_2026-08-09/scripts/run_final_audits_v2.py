from __future__ import annotations

"""执行第二轮科学、引用、非复制、版本、文档与期刊合规审计。"""

import hashlib
import json
import re
import unicodedata
from collections import defaultdict
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document
from pypdf import PdfReader


WORKSPACE = Path(__file__).resolve().parents[3]
ROOT = Path(__file__).resolve().parents[1]
QA = ROOT / "qa"
MANUSCRIPT_MD = ROOT / "documents" / "EAAI_manuscript_source_v2.md"
SUPP_MD = ROOT / "documents" / "EAAI_supplementary_source_v2.md"
FINAL_RESULTS = WORKSPACE / "paper_runs" / "multimap_v3_2_14" / "formal_evaluation" / "results" / "final_results.jsonl"
FINAL_AUDIT = WORKSPACE / "paper_runs" / "multimap_v3_2_14" / "formal_evaluation" / "results" / "final_audit_status.json"
PROTOCOL = WORKSPACE / "paper_runs" / "protocols" / "multimap_generalization_v3_2_14" / "protocol.json"
V1 = WORKSPACE / "paper_delivery" / "EAAI_2026-08-09"

EXPECTED = {
    "results_sha256": "4b620c21566c2e33c875f6bea2017b741b02a7d30d70aa50add60a6d06214a2c",
    "protocol_sha256": "c0ac70fb8fac32bd60afe602b24d8534d9329e531ea381d9a73bf4237c9fbc58",
    "v1_tree_recorded": "8703b1fc9bfdb75791fe094b8e3ef2ca8a726de1f1d0bf28bf0b021b02e41d80",
    "v1_files": 227,
    "v2_snapshot_time": "2026-08-09T13:15:35.1024338+08:00",
}


def sha(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def norm(text: str) -> str:
    text = unicodedata.normalize("NFKD", text)
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    return re.sub(r"[^a-z0-9]+", " ", text.lower()).strip()


def words(text: str) -> list[str]:
    return norm(text).split()


def sentences(text: str) -> list[str]:
    text = re.sub(r"^#+.*$", " ", text, flags=re.M)
    text = re.sub(r"\([^)]*\d{4}[^)]*\)", " ", text)
    parts = re.split(r"(?<=[.!?])\s+(?=[A-Z])", text)
    return [p.strip() for p in parts if len(words(p)) >= 10]


def scientific_audit() -> dict:
    audit = json.loads(FINAL_AUDIT.read_text(encoding="utf-8-sig"))
    row_count = sum(1 for _ in FINAL_RESULTS.open("rb"))
    text = MANUSCRIPT_MD.read_text(encoding="utf-8")
    checks = {
        "final_audit_passed": audit.get("passed") is True,
        "rows_21648": row_count == 21648 == audit.get("row_count") == audit.get("route_count"),
        "results_sha256": sha(FINAL_RESULTS) == EXPECTED["results_sha256"] == audit.get("results_sha256"),
        "protocol_sha256": sha(PROTOCOL) == EXPECTED["protocol_sha256"],
        "ppo_mlp_absent": audit.get("ppo_mlp_absent") is True and "ppo_mlp" not in text.lower(),
        "traditional_ppo_present": audit.get("traditional_ppo_present") is True,
        "map_is_independent_unit": "The map was the independent unit" in text,
        "confirmatory_endpoint_named": "confirmatory endpoint was safe weighted coverage" in text.lower(),
        "bounded_claims_present": all(x in text.lower() for x in ["not universal optimality", "simulation rather than flight validation", "did not test 28-, 32- or larger-point extrapolation"]),
    }
    return {"status": "PASS" if all(checks.values()) else "FAIL", "checks": checks, "row_count": row_count,
            "family_counts": audit.get("family_counts")}


def exemplar_audit() -> dict:
    matrix = json.loads((ROOT / "literature" / "eaai_12_fulltext_matrix_source.json").read_text(encoding="utf-8"))
    records = matrix.get("papers", matrix if isinstance(matrix, list) else [])
    pdf_dir = ROOT / "literature" / "EAAI_exemplar_fulltexts_v2"
    rows = []
    for rec in records:
        number = int(rec.get("number", rec.get("no", rec.get("id", 0))))
        pdf = pdf_dir / f"{number}.pdf"
        pages = len(PdfReader(str(pdf)).pages) if pdf.exists() else 0
        rows.append({"number": number, "exists": pdf.exists(), "sha256_match": pdf.exists() and sha(pdf) == rec["sha256"],
                     "pages_match": pages == int(rec["pages"]), "pages": pages, "doi": rec["doi"]})
    ok = len(rows) == 12 and all(r["exists"] and r["sha256_match"] and r["pages_match"] for r in rows)
    return {"status": "PASS" if ok else "FAIL", "count": len(rows), "records": rows}


def citation_audit() -> dict:
    text = MANUSCRIPT_MD.read_text(encoding="utf-8")
    body, ref_text = text.split("## References", 1)
    blocks = [b.strip() for b in re.split(r"\n\s*\n", ref_text) if b.strip()]
    selection = json.loads((ROOT / "literature" / "reference_selection_v2.json").read_text(encoding="utf-8"))
    final_refs = [r for r in selection["references"] if r["status"] == "final"]
    uncited = []
    for block in blocks:
        lead = block.split(",", 1)[0].strip()
        if norm(lead) not in norm(body):
            uncited.append(lead)
    dois = sorted(set(re.findall(r"10\.\d{4,9}/[-._;()/:A-Za-z0-9]+", ref_text)))
    ris = (ROOT / "literature" / "verified_references_v2.ris").read_text(encoding="utf-8")
    ris_records = len(re.findall(r"^TY  - ", ris, flags=re.M))
    status = json.loads((ROOT / "literature" / "reference_status_v2.json").read_text(encoding="utf-8"))
    checks = {
        "no_citation_needed_placeholders": "[CITATION NEEDED" not in text,
        "reference_count_30": len(blocks) == len(final_refs) == 30,
        "ris_count_30": ris_records == 30,
        "fulltext_status_count_30": int(status.get("reference_count", 0)) == 30,
        "all_references_cited": len(uncited) == 0,
        "school_access_list_empty": "No school-access request is currently required" in (ROOT / "literature" / "School_Access_Required_List.md").read_text(encoding="utf-8"),
    }
    return {"status": "PASS" if all(checks.values()) else "FAIL", "checks": checks, "uncited": uncited,
            "doi_count": len(dois), "final_reference_count": len(final_refs)}


def noncopy_audit() -> dict:
    manuscript = MANUSCRIPT_MD.read_text(encoding="utf-8").split("## References", 1)[0]
    ms = sentences(manuscript)
    ex_dir = WORKSPACE / "tmp" / "pdfs" / "eaai_exemplars"
    exemplar_sentences = []
    for p in sorted(ex_dir.glob("*.txt"), key=lambda x: int(x.stem)):
        for s in sentences(p.read_text(encoding="utf-8", errors="ignore")):
            exemplar_sentences.append((p.stem, s, words(s)))

    index: dict[tuple[str, ...], list[int]] = defaultdict(list)
    for idx, (_, _, ws) in enumerate(exemplar_sentences):
        for i in range(max(0, len(ws) - 7)):
            index[tuple(ws[i:i + 8])].append(idx)
    matches = []
    for sentence in ms:
        ws = words(sentence)
        candidates = set()
        for i in range(max(0, len(ws) - 7)):
            candidates.update(index.get(tuple(ws[i:i + 8]), []))
        for idx in candidates:
            paper, ex_sentence, ex_words = exemplar_sentences[idx]
            ratio = SequenceMatcher(None, " ".join(ws), " ".join(ex_words)).ratio()
            if ratio >= 0.55:
                matches.append({"ratio": round(ratio, 4), "paper": paper, "manuscript": sentence, "exemplar": ex_sentence})
    matches.sort(key=lambda x: x["ratio"], reverse=True)
    high = [m for m in matches if m["ratio"] >= 0.78]
    return {"status": "PASS" if not high else "REVIEW", "method": "Local sentence comparison after citation removal; exact 8-word shingle candidate generation plus normalized SequenceMatcher", "threshold": 0.78,
            "manuscript_sentences": len(ms), "exemplar_sentences": len(exemplar_sentences), "high_similarity_count": len(high), "top_matches": matches[:20]}


def document_and_journal_audit() -> dict:
    main_docx = ROOT / "deliverables" / "EAAI_manuscript_anonymized_v2.docx"
    doc = Document(main_docx)
    doc_text = "\n".join(p.text for p in doc.paragraphs)
    md = MANUSCRIPT_MD.read_text(encoding="utf-8")
    abstract = md.split("## Abstract", 1)[1].split("**Keywords:**", 1)[0]
    abstract_words = len(re.findall(r"\b[\w–-]+\b", abstract))
    keyword_line = re.search(r"\*\*Keywords:\*\*\s*(.+)", md).group(1)
    keyword_count = len([x for x in keyword_line.split(";") if x.strip()])
    highlights = [line[2:].strip() for line in (ROOT / "documents" / "EAAI_highlights_source_v2.md").read_text(encoding="utf-8").splitlines() if line.startswith("- ")]
    pdfs = {}
    for stem in ["EAAI_manuscript_anonymized_v2", "EAAI_title_page_v2", "EAAI_supplementary_material_v2", "EAAI_highlights_v2", "EAAI_cover_letter_v2"]:
        pdf = ROOT / "qa" / "docx_render" / stem / f"{stem}.pdf"
        pdfs[stem] = len(PdfReader(str(pdf)).pages)
    author_queue = (ROOT / "evidence" / "Author_Verification_Queue_v2.md").read_text(encoding="utf-8")
    checks = {
        "abstract_le_250": abstract_words <= 250,
        "keywords_1_to_6": 1 <= keyword_count <= 6,
        "highlights_3_to_5": 3 <= len(highlights) <= 5,
        "highlights_le_85_chars": all(len(h) <= 85 for h in highlights),
        "main_pages_le_50": pdfs["EAAI_manuscript_anonymized_v2"] <= 50,
        "five_docx_present": len(list((ROOT / "deliverables").glob("EAAI_*_v2.docx"))) == 5,
        "single_column": all(not sec._sectPr.xpath("./w:cols[@w:num and @w:num!='1']") for sec in doc.sections),
        "anonymous_main": all(x.lower() not in doc_text.lower() for x in ["C:\\Users\\xsp", "DRL代码", "@gmail.com", "@qq.com"]),
        "figure_captions_8": sum(1 for p in doc.paragraphs if re.match(r"Figure [1-8]\. ", p.text)) == 8,
        "required_declarations_present": all(x in doc_text for x in ["Data and code availability", "Declaration of competing interest", "Funding", "Declaration of generative AI"]),
        "all_author_placeholders_queued": "AQ01" in author_queue and "AQ09" in author_queue,
        "no_citation_needed": "[CITATION NEEDED" not in doc_text,
    }
    return {"status": "PASS_WITH_AUTHOR_INPUT" if all(checks.values()) else "FAIL", "checks": checks,
            "abstract_words": abstract_words, "keyword_count": keyword_count, "highlight_lengths": [len(h) for h in highlights], "pdf_pages": pdfs,
            "render_backend": "Microsoft Word COM export followed by Poppler page rendering; LibreOffice unavailable"}


def version_audit() -> dict:
    files = [p for p in V1.rglob("*") if p.is_file()]
    latest = max(p.stat().st_mtime for p in files)
    import datetime
    latest_iso = datetime.datetime.fromtimestamp(latest, datetime.timezone.utc).astimezone().isoformat()
    snap = datetime.datetime.fromisoformat(EXPECTED["v2_snapshot_time"])
    latest_dt = datetime.datetime.fromtimestamp(latest, snap.tzinfo)
    # 独立、可复算的当前指纹；基线记录采用的串联格式未单独序列化，因此同时使用文件数和修改时间闸门。
    rows = []
    for p in sorted(files, key=lambda x: x.relative_to(V1).as_posix()):
        rows.append(f"{p.relative_to(V1).as_posix()}|{sha(p)}|{p.stat().st_size}")
    current_fingerprint = hashlib.sha256("\n".join(rows).encode("utf-8")).hexdigest()
    checks = {"file_count_unchanged": len(files) == EXPECTED["v1_files"], "no_file_modified_after_v2_snapshot": latest_dt <= snap,
              "v1_recorded_hash_retained": EXPECTED["v1_tree_recorded"] in (ROOT / "evidence" / "V2_Preflight_Report.md").read_text(encoding="utf-8")}
    return {"status": "PASS" if all(checks.values()) else "FAIL", "checks": checks, "recorded_preflight_tree_sha256": EXPECTED["v1_tree_recorded"],
            "current_independent_fingerprint": current_fingerprint, "current_fingerprint_rule": "SHA256 of sorted relative_path|file_sha256|bytes rows", "latest_v1_mtime": latest_iso}


def write_report(name: str, data: dict) -> None:
    (QA / f"{name}_v2.json").write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = [f"# {name.replace('_', ' ').title()} v2", "", f"- Status: **{data['status']}**", ""]
    if "checks" in data:
        lines.append("## Checks\n")
        for k, v in data["checks"].items():
            lines.append(f"- {'PASS' if v else 'FAIL'} — `{k}`")
    lines.extend(["", "## Machine-readable detail", "", f"See `{name}_v2.json`."])
    (QA / f"{name}_v2.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    QA.mkdir(parents=True, exist_ok=True)
    audits = {
        "Scientific_Audit": scientific_audit(),
        "Exemplar_Fulltext_Audit": exemplar_audit(),
        "Citation_Audit": citation_audit(),
        "Noncopy_Audit": noncopy_audit(),
        "Document_and_EAAI_Compliance_Audit": document_and_journal_audit(),
        "Version_Isolation_Audit": version_audit(),
        "Binary_Delivery_Validation": json.loads((QA / "Binary_Delivery_Validation_v2.json").read_text(encoding="utf-8")),
        "Code_and_Workspace_Validation": json.loads((QA / "Code_and_Workspace_Validation_v2.json").read_text(encoding="utf-8")),
    }
    for name, data in audits.items():
        write_report(name, data)
    overall = all(v["status"] in {"PASS", "PASS_WITH_AUTHOR_INPUT"} for v in audits.values())
    summary = {"status": "PASS_WITH_AUTHOR_INPUT" if overall else "REVIEW_REQUIRED", "audits": {k: v["status"] for k, v in audits.items()},
               "author_side_items": "See evidence/Author_Verification_Queue_v2.md", "nature_skills_used": False, "experiments_rerun": False}
    (QA / "Final_Compliance_Report_v2.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    md = ["# Final Compliance Report v2", "", f"Overall status: **{summary['status']}**", "", "## Audit results", ""]
    md.extend(f"- {k}: **{v}**" for k, v in summary["audits"].items())
    md += ["", "## Remaining author-side verification", "", "All unknown author, affiliation, funding, competing-interest, CRediT, repository and disclosure approvals remain in `evidence/Author_Verification_Queue_v2.md`; no value was guessed.", "", "No Nature skill, generative-image tool, model retraining, protocol change or result recomputation was used."]
    (QA / "Final_Compliance_Report_v2.md").write_text("\n".join(md) + "\n", encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
