from __future__ import annotations

"""从第二轮 Markdown 源稿生成 EAAI 单栏、双盲 DOCX 交付包。"""

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "documents"
FIGS = ROOT / "figures" / "submission"
OUT = ROOT / "deliverables"
QA = ROOT / "qa"

# 重要可调参数：与 documents skill 的 narrative_proposal 预设一致，
# 并采用科学论文覆盖项（Times New Roman、Letter、单栏、1 inch 页边距）。
BODY_PT = 10.5
LINE_MULTIPLE = 1.15
FIG_WIDTH = 6.35


MAIN_FIGURES = {
    "[FIGURE 1": [FIGS / "main" / "F01_method_and_evaluation_workflow.png"],
    "[FIGURE 2": [FIGS / "main" / "M01_priority_weighted_coverage_english.png", FIGS / "main" / "M03_priority_stratum_effects_english.png"],
    "[FIGURE 3": [FIGS / "main" / "M02_safety_and_return_effects_english.png", FIGS / "main" / "M04_resource_use_english.png"],
    "[FIGURE 4": [FIGS / "main" / "M05_online_planning_time_ECDF_repaired.png", FIGS / "supplementary" / "S02_quality_time_tradeoff_english.png"],
    "[FIGURE 5": [FIGS / "main" / "M06_training_curves_english.png", FIGS / "main" / "M07_training_stability_efficiency_english.png"],
    "[FIGURE 6": [FIGS / "main" / "M08_unseen_maps_and_DSM_transfer_english.png", FIGS / "showcase" / "V02_fixed_DSM_route_repaired.png"],
    "[FIGURE 7": [FIGS / "main" / "M09_two_layer_robustness_english.png"],
    "[FIGURE 8": [FIGS / "main" / "M10_ablation_effects_english.png"],
}

CAPTIONS = {
    "[FIGURE 1": "Figure 1. Return-aware PPO–Pointer planning and frozen evaluation workflow. The diagram describes the implemented sequence from mountain-road task representation to composite return-feasibility screening and the predefined evidence outputs; it is a non-generative schematic and does not add experimental evidence.",
    "[FIGURE 2": "Figure 2. Coverage and priority-stratum evidence. (a) Map-level safe priority-weighted coverage on 24 unseen synthetic maps and eight DSM maps; points are independent-map aggregates and vertical markers are medians. (b) Full-model minus ablation differences across priority strata; these stratified values are descriptive and do not replace the confirmatory map-level endpoint.",
    "[FIGURE 3": "Figure 3. Safe-return and resource outcomes. (a) PPO–Pointer minus comparator percentage-point effects for safe completion and depot return under known shifts and hidden model/perception mismatch; intervals are map-level 95% bootstrap intervals. (b) Median energy, range and mission-time utilization among safe routes; the dashed line denotes the budget limit.",
    "[FIGURE 4": "Figure 4. Online planning time and quality–time trade-off. (a) Empirical cumulative distributions of per-task online planning time over the frozen evaluation jobs. (b) Map-level D1 versus 95th-percentile planning time on synthetic and DSM tasks. Timing values are specific to the frozen software and hardware protocol and are not cross-platform latency guarantees.",
    "[FIGURE 5": "Figure 5. Training trajectories, stability and sample efficiency. (a) Five training seeds per principal learner; thin curves are seeds, heavy curves are medians and bands are interquartile ranges. (b) Post-hoc D6 stability and D7 sample-efficiency scores. These training dimensions supplement rather than redefine the confirmatory final-route endpoint.",
    "[FIGURE 6": "Figure 6. Cross-map performance and zero-shot DSM simulation transfer. (a) Map-level D1 estimates and 95% bootstrap intervals on unseen procedural maps and eight Copernicus DSM maps. (b) Terrain, road network, inspection points and representative routes for one fixed DSM task. The route panel is illustrative and non-inferential; DSM evaluation remains simulation, not flight validation.",
    "[FIGURE 7": "Figure 7. Two-layer robustness evaluation. D1 retention is reported separately for shifts revealed to the planner and hidden model/perception mismatches. The dashed reference denotes retention of 1.0; maps, not tasks or routes, are the inferential units.",
    "[FIGURE 8": "Figure 8. Four component ablations. Points are full-model minus ablation mean differences in map-level safe weighted coverage; horizontal intervals are 95% map-bootstrap intervals. Asterisked significance, where reported in the Source Data, follows Holm adjustment within the prespecified ablation family.",
}

SUPP_FIGURES = [
    ("Figure S1", FIGS / "supplementary" / "S01_performance_profile_english.png", "Regret performance profile for all frozen algorithms. Each curve is an empirical cumulative distribution over the stated task set; the panel is descriptive and does not change the map-level inferential unit."),
    ("Figure S2", FIGS / "supplementary" / "S02_quality_time_tradeoff_english.png", "Coverage–time trade-off on synthetic and DSM tasks using D1 and the protocol-specific 95th-percentile planning time."),
    ("Figure S3", FIGS / "supplementary" / "S03_oracle_regret_cost_english.png", "Traditional-planner oracle-regret intervals versus planning time. Solver status and certification share must be read together with these values."),
    ("Figure S4", FIGS / "supplementary" / "S04_scenario_heatmap_english.png", "Scenario-stratified mean D1 for the three principal learning models. Scenario cells are descriptive strata, not independent replication units."),
    ("Figure S5", FIGS / "supplementary" / "S05_failure_modes_english.png", "Robustness and failure-mode rates under known shifts and hidden mismatch. Values are frozen task aggregates; safe and return rates are not flight-certification evidence."),
    ("Figure S6", FIGS / "supplementary" / "S06_seven_model_training_english.png", "Training trajectories for the seven paper-eligible learning variants. Curves summarize five seeds per model over 3,000 episodes."),
    ("Figure S7", FIGS / "supplementary" / "S07_posthoc_score_english.png", "Seven normalized dimensions and the post-hoc 100-point composite. This score is exploratory, weight-dependent and is not the confirmatory endpoint."),
    ("Figure S8", FIGS / "supplementary" / "S08_weight_sensitivity_english.png", "Joint sensitivity of the first-place share to training weight and operational floor in the post-hoc composite analysis."),
    ("Figure S9", FIGS / "showcase" / "V01_fixed_synthetic_route_english.png", "Representative routes on one fixed synthetic task. The panel is illustrative and not used for statistical inference."),
    ("Figure S10", FIGS / "showcase" / "V02_fixed_DSM_route_repaired.png", "Representative routes on one fixed DSM task with elevation and road context. This is a zero-shot simulation-transfer illustration, not field validation."),
]


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])


def configure_document(doc: Document, running_header: str) -> None:
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin = sec.bottom_margin = Inches(0.82)
    sec.left_margin = sec.right_margin = Inches(0.92)
    sec.header_distance = sec.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"; normal.font.size = Pt(BODY_PT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = LINE_MULTIPLE
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.widow_control = True

    for name, size, color in [("Title", 17, "183A5A"), ("Heading 1", 13.5, "183A5A"), ("Heading 2", 12, "315A7A"), ("Heading 3", 10.8, "3D6078")]:
        st = styles[name]; st.font.name = "Times New Roman"; st.font.size = Pt(size); st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True; st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.paragraph_format.keep_with_next = True; st.paragraph_format.space_before = Pt(9); st.paragraph_format.space_after = Pt(4)

    if "Figure Caption" not in styles:
        fc = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        fc = styles["Figure Caption"]
    fc.font.name = "Times New Roman"; fc.font.size = Pt(8.5); fc.font.color.rgb = RGBColor(45, 45, 45)
    fc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; fc.paragraph_format.space_after = Pt(8); fc.paragraph_format.keep_with_next = True

    header = sec.header.paragraphs[0]
    header.text = running_header; header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = "Times New Roman"; header.runs[0].font.size = Pt(8); header.runs[0].font.color.rgb = RGBColor(95, 105, 115)
    add_page_field(sec.footer.paragraphs[0])


def add_inline(paragraph, text: str) -> None:
    """处理有限 Markdown 行内格式，避免把控制标记留在最终 DOCX。"""
    token = re.compile(r"(\*\*.*?\*\*|`.*?`|\*.*?\*)")
    pos = 0
    for match in token.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        value = match.group(0)
        if value.startswith("**"):
            r = paragraph.add_run(value[2:-2]); r.bold = True
        elif value.startswith("`"):
            r = paragraph.add_run(value[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(8.5)
        else:
            r = paragraph.add_run(value[1:-1]); r.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def sorted_reference_lines(lines: list[str]) -> list[str]:
    try:
        idx = lines.index("## References")
    except ValueError:
        return lines
    head = lines[:idx + 1]
    text = "\n".join(lines[idx + 1:]).strip()
    blocks = [b.strip() for b in re.split(r"\n\s*\n", text) if b.strip()]
    blocks.sort(key=lambda s: re.sub(r"[^A-Za-z]", "", s).lower())
    tail: list[str] = []
    for b in blocks:
        tail.extend(b.splitlines()); tail.append("")
    return head + [""] + tail


def add_table(doc: Document, rows: list[list[str]]) -> None:
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.cell(i, j); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, row[j].strip() if j < len(row) else "")
            for run in p.runs:
                run.font.name = "Times New Roman"; run.font.size = Pt(8.1); run.bold = i == 0
        if i == 0:
            set_repeat_table_header(table.rows[i])
            for cell in table.rows[i].cells:
                shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "DCE8F2"); cell._tc.get_or_add_tcPr().append(shd)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_picture(doc: Document, path: Path, width: float = FIG_WIDTH) -> None:
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))


def add_main_figure(doc: Document, marker: str) -> None:
    key = next(k for k in MAIN_FIGURES if marker.startswith(k))
    for path in MAIN_FIGURES[key]:
        add_picture(doc, path, FIG_WIDTH if len(MAIN_FIGURES[key]) == 1 else 6.05)
    p = doc.add_paragraph(style="Figure Caption"); add_inline(p, CAPTIONS[key])


def markdown_to_doc(doc: Document, path: Path, main: bool = False) -> None:
    lines = path.read_text(encoding="utf-8").splitlines()
    if main:
        lines = sorted_reference_lines(lines)
    i = 0
    first_title = True
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1; continue
        if main and any(line.startswith(k) for k in MAIN_FIGURES):
            add_main_figure(doc, line); i += 1; continue
        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r"\s*:?-+:?\s*", c) for c in cells):
                    rows.append(cells)
                i += 1
            if rows: add_table(doc, rows)
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Title" if first_title else "Heading 1"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER if first_title else WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, line[2:]); first_title = False; i += 1; continue
        if line.startswith("## "):
            if line == "## References":
                doc.add_page_break()
            p = doc.add_paragraph(style="Heading 1"); add_inline(p, line[3:]); i += 1; continue
        if line.startswith("### "):
            style = "Figure Caption" if re.match(r"### Table", line) else "Heading 2"
            p = doc.add_paragraph(style=style); add_inline(p, line[4:]); i += 1; continue
        if line.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3"); add_inline(p, line[5:]); i += 1; continue
        if re.match(r"^[-*] ", line):
            p = doc.add_paragraph(style="List Bullet"); add_inline(p, line[2:]); i += 1; continue
        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number"); add_inline(p, re.sub(r"^\d+\. ", "", line)); i += 1; continue

        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r"^(#{1,4} |[-*] |\d+\. |\|)", lines[i].strip()) and not (main and any(lines[i].startswith(k) for k in MAIN_FIGURES)):
            para_lines.append(lines[i].strip()); i += 1
        p = doc.add_paragraph(); add_inline(p, " ".join(para_lines))


def save_doc(md_name: str, out_name: str, header: str, main: bool = False, supplement: bool = False) -> Path:
    doc = Document(); configure_document(doc, header)
    markdown_to_doc(doc, DOCS / md_name, main=main)
    if supplement:
        doc.add_page_break(); doc.add_heading("Supplementary figures", level=1)
        for title, path, caption in SUPP_FIGURES:
            add_picture(doc, path, 6.1)
            p = doc.add_paragraph(style="Figure Caption"); add_inline(p, f"{title}. {caption}")
    out = OUT / out_name; out.parent.mkdir(parents=True, exist_ok=True); doc.save(out)
    return out


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True); QA.mkdir(parents=True, exist_ok=True)
    produced = [
        save_doc("EAAI_manuscript_source_v2.md", "EAAI_manuscript_anonymized_v2.docx", "EAAI Original Research | Anonymized manuscript", main=True),
        save_doc("EAAI_title_page_source_v2.md", "EAAI_title_page_v2.docx", "EAAI Original Research | Title page"),
        save_doc("EAAI_supplementary_source_v2.md", "EAAI_supplementary_material_v2.docx", "EAAI Original Research | Supplementary material", supplement=True),
        save_doc("EAAI_highlights_source_v2.md", "EAAI_highlights_v2.docx", "EAAI Original Research | Highlights"),
        save_doc("EAAI_cover_letter_source_v2.md", "EAAI_cover_letter_v2.docx", "EAAI Original Research | Cover letter"),
    ]
    report = {
        "preset": "narrative_proposal",
        "named_override": "EAAI scientific manuscript: Times New Roman, Letter, single column, 0.92-inch side margins",
        "body_pt": BODY_PT,
        "line_multiple": LINE_MULTIPLE,
        "documents": [str(p.relative_to(ROOT)) for p in produced],
    }
    (QA / "docx_build_profile_v2.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
