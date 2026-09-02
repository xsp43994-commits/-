from __future__ import annotations

import csv
import hashlib
import json
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image


# 关键参数：所有稿件均使用 EAAI 单栏匿名投稿布局；输入只读，输出集中到交付目录。
WORKSPACE = Path(r"C:\Users\xsp\Desktop\DRL代码")
DELIVERY = WORKSPACE / "paper_delivery" / "EAAI_2026-08-09"
DOCS = DELIVERY / "documents"
FIG_MAIN = DELIVERY / "figures" / "main"
FIG_SUPP = DELIVERY / "figures" / "supplementary"
FIG_DOCX = DELIVERY / "figures" / "docx_embedded_previews"
EVIDENCE = DELIVERY / "evidence"
FROZEN = WORKSPACE / "paper_runs" / "multimap_v3_2_14"
STAT = FROZEN / "analysis" / "pre_plot_statistics"
FIG_SOURCE = FROZEN / "figures" / "paper_redraw_multibackend_v3" / "source_data"
REGISTER = json.loads((DELIVERY / "literature" / "literature_register.json").read_text(encoding="utf-8"))

TITLE = "Return-aware proximal policy optimization with pointer attention for multi-resource mountain-road unmanned aerial vehicle inspection planning"
SHORT_TITLE = "Return-aware learning for mountain-road inspection planning"
KEYWORDS = ["unmanned aerial vehicle inspection", "reinforcement learning", "combinatorial planning", "resource constraints", "geographic simulation transfer"]

HIGHLIGHTS = [
    "Return-aware masking enforces energy, distance, time, and depot feasibility",
    "Pointer attention supports variable-size fixed-point inspection decisions",
    "PPO and A2C attain similar safe coverage across unseen map simulations",
    "PPO improves training stability and sample efficiency over A2C",
    "Classical solvers expose an explicit coverage-versus-online-time trade-off",
]

for item in HIGHLIGHTS:
    if len(item) > 85:
        raise ValueError(f"Highlight exceeds 85 characters: {len(item)} {item}")


REFERENCE_ORDER = [
    "10.1016/j.engappai.2026.115219",
    "10.1016/j.engappai.2025.112090",
    "10.1016/j.engappai.2024.109870",
    "10.1016/j.engappai.2024.109339",
    "10.1016/j.engappai.2023.105891",
    "10.1016/j.engappai.2023.106703",
    "10.1016/j.engappai.2022.105321",
    "10.1016/j.engappai.2022.105182",
    "10.1109/twc.2019.2902559",
    "10.1109/iros.2017.8202133",
    "10.1007/s10994-021-05961-4",
    "Schulman2017PPO",
    "Vinyals2015Pointer",
    "Vaswani2017Attention",
    "Schulman2016GAE",
    "Mnih2016A3C",
    "Kool2019AttentionRouting",
    "Achiam2017CPO",
    "Garcia2015SafeRL",
    "CopernicusDEM2021",
    "10.1080/01621459.1937.10503522",
    "10.2307/3001968",
    "10.1214/aoms/1177704172",
    "10.1214/aos/1176344552",
    "Demsar2006",
    "Holm1979",
    "10.1016/j.autcon.2024.105764",
]

REFERENCE_LOOKUP = {(x.get("doi") or x.get("id", "")).lower(): x for x in REGISTER["references"]}
REFERENCES = [REFERENCE_LOOKUP[x.lower()] for x in REFERENCE_ORDER]


CAPTIONS = {
    "Fig01_method": "Return-aware proximal policy optimization with pointer attention and the locked evidence workflow. Candidate inspection points are encoded with priority-aware self-attention. The pointer policy scores variable-length candidates, while a deterministic feasibility mask excludes actions that cannot satisfy visit, dynamic, resource, and depot-return conditions. Training and all evaluations used the frozen v3.2.14 protocol.",
    "Fig02_coverage_priority": "Coverage and priority effects on unseen synthetic maps and independent geographic digital surface model (DSM) simulations. (a) Map-level safe priority-weighted coverage. Points denote independent maps and horizontal bars denote medians. (b) Priority-stratum coverage for PPO+Pointer and the no-priority-bias ablation; these stratum contrasts are descriptive.",
    "Fig03_safety_resources": "Safety/return effects and resource use. (a) PPO+Pointer minus comparator percentage-point effects under known shifts and hidden mismatch. Intervals are the frozen map-level uncertainty estimates. (b) Median energy, distance, and mission-time budget utilisation among safe routes.",
    "Fig04_time_tradeoff": "Online planning time and quality–time trade-off. (a) Empirical cumulative distributions of planning time. (b) Safe priority-weighted coverage dimension versus 95th-percentile planning time; marker size encodes safe-route share.",
    "Fig05_training": "Training trajectories, stability, and sample efficiency. (a) Median safe priority-weighted coverage over five seeds with interquartile bands. (b) Frozen D6 training-stability and D7 sample-efficiency scores.",
    "Fig06_transfer": "Unseen-map performance and zero-shot geographic DSM simulation transfer. (a) Map-level estimates and uncertainty intervals on unseen procedural maps and eight independent Copernicus DSM regions. (b) Representative Taihang DSM simulation with depot, inspection points, and routes. This panel is a simulation result, not a flight experiment.",
    "Fig07_robustness": "Two-layer robustness. Cells show retention relative to each model's own reference performance; S denotes safe-route share. Known shifts were represented in the training randomisation envelope, whereas hidden mismatch conditions were deliberately withheld.",
    "Fig08_ablation": "Four-component ablation. Points are Hodges–Lehmann estimates for PPO+Pointer minus each ablation; intervals are frozen 95% hierarchical-bootstrap intervals. The other three individual removals were not significant after Holm adjustment.",
}


FACT_MAP = [
    ("F01", "The frozen protocol identity is multimap_v3_2_14 (v3.2.14).", "HANDOFF.md; final_audit_status.json", "metadata", "verified"),
    ("F02", "The final result file contains 21,648 records and excludes ppo_mlp.", "final_audit_status.json; final_results.jsonl", "result", "verified"),
    ("F03", "Training used 72 procedural mountain maps, 648 tasks, five seeds, and 3,000 episodes for each of seven learning variants.", "HANDOFF.md; training summaries", "method", "verified"),
    ("F04", "Formal evaluation used 24 unseen procedural maps (216 tasks) and eight independent Copernicus DSM regions (144 tasks).", "evaluation matrix; HANDOFF.md", "method", "verified"),
    ("F05", "Node counts 16, 20, and 24 all occurred in training; evaluation does not test out-of-range scale generalisation.", "protocol; HANDOFF.md", "limitation", "verified"),
    ("F06", "The full actor–critic contains 341,505 parameters for the audited seed-42 checkpoint.", "checkpoint audit", "method", "verified"),
    ("F07", "The traditional comparator is a fixed-slot FlatMLPActorCritic without attention.", "implementation; HANDOFF.md", "method", "verified"),
    ("F08", "Safe priority-weighted coverage assigns zero to unsafe or hard-violation routes.", "protocol and analysis code", "method", "verified"),
    ("F09", "Maps are the independent statistical unit; task, seed, route, and road repetitions are not treated as independent maps.", "statistics protocol", "statistic", "verified"),
    ("F10", "PPO+Pointer mean safe weighted coverage is 0.4858 on unseen synthetic maps and 0.5024 on DSM simulations.", "descriptive_metrics.csv", "result", "verified"),
    ("F11", "PPO+Pointer and A2C+Pointer do not differ significantly on the confirmatory safe coverage outcome.", "confirmatory_pairwise.csv", "result", "verified"),
    ("F12", "ACO, SA, and MILP can exceed PPO+Pointer on safe weighted coverage, with higher online computation.", "descriptive_metrics.csv; M05/S02 Source Data", "result", "verified"),
    ("F13", "Removing the return reserve produces a large, significant coverage loss; the other three individual ablations do not.", "confirmatory_pairwise.csv", "result", "verified"),
    ("F14", "Geographic DSM results are zero-shot simulations and are not real-flight validation or safety certification.", "protocol; HANDOFF.md", "limitation", "verified"),
    ("F15", "The post-hoc 100-point operational score is descriptive and not a confirmatory primary outcome.", "operational analysis manifest", "limitation", "verified"),
]

CLAIM_MAP = [
    ("C01", "The method enforces return-aware resource feasibility before action selection.", "F08; Fig. 1; implementation", "method fact", "show", "No formal safety certification"),
    ("C02", "Pointer attention supports variable-length candidate selection within the trained node-count range.", "F05; Fig. 1", "method fact", "supports", "No claim beyond 16/20/24 nodes"),
    ("C03", "PPO+Pointer strongly improves safe weighted coverage over Flat-MLP PPO.", "F10; Fig. 2; pairwise statistics", "result", "shows", "Simulation domains only"),
    ("C04", "PPO+Pointer and A2C+Pointer attain statistically similar confirmatory safe coverage.", "F11; Fig. 2", "result", "indicates", "Non-significance is not proof of equivalence"),
    ("C05", "PPO+Pointer has a stronger combined training profile than A2C+Pointer.", "Fig. 5; D6/D7 frozen scores", "result", "indicates", "Mechanism not causally isolated"),
    ("C06", "Return reserve is the dominant independently tested safety-related component.", "F13; Fig. 8", "result", "supports", "Other submasks were not separately ablated"),
    ("C07", "Classical solvers expose a quality–time trade-off rather than universal PPO dominance.", "F12; Fig. 4", "result", "shows", "Hardware-specific time values"),
    ("C08", "The policy transfers zero-shot to independent geographic DSM simulations.", "F04; F14; Fig. 6", "result", "shows", "Not physical-domain transfer"),
    ("C09", "Hidden mismatch degrades performance but preserves a PPO advantage over Flat-MLP PPO.", "Fig. 7; robustness tables", "result", "indicates", "Finite perturbation catalogue"),
    ("C10", "The study provides reproducible simulation evidence for engineering planning decisions.", "F01–F15; package manifest", "synthesis", "supports", "Requires independent external and flight validation"),
]


def read_csv(name: str) -> list[dict[str, str]]:
    with (STAT / name).open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


DESCRIPTIVE = read_csv("descriptive_metrics.csv")
PAIRWISE = read_csv("confirmatory_pairwise.csv")
OMNIBUS = read_csv("confirmatory_omnibus.csv")


def find_row(rows: list[dict[str, str]], **conditions: str) -> dict[str, str]:
    matches = [row for row in rows if all(row.get(key) == value for key, value in conditions.items())]
    if len(matches) != 1:
        raise KeyError(f"Expected one row for {conditions}, found {len(matches)}")
    return matches[0]


def f3(value: str | float) -> str:
    return f"{float(value):.3f}"


def sci(value: str | float) -> str:
    return f"{float(value):.3g}"


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None, size: float = 8.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def new_document(title: str, anonymized: bool = False) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.header_distance = Cm(0.8)
    sec.footer_distance = Cm(0.8)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for name, size, color in [("Title", 16, "17365D"), ("Heading 1", 13, "17365D"), ("Heading 2", 11.5, "2F5597"), ("Heading 3", 10.5, "2F5597")]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(4)
    if "Caption EAAI" not in styles:
        cap = styles.add_style("Caption EAAI", WD_STYLE_TYPE.PARAGRAPH)
        cap.font.name = "Arial"
        cap.font.size = Pt(8.5)
        cap.paragraph_format.space_after = Pt(8)
        cap.paragraph_format.keep_with_next = False
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    header = sec.header.paragraphs[0]
    header.text = "Anonymous manuscript" if anonymized else SHORT_TITLE
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        run.font.name = "Arial"; run.font.size = Pt(8); run.font.color.rgb = RGBColor(100, 116, 139)
    set_page_number(sec.footer.paragraphs[0])
    doc.core_properties.title = title
    doc.core_properties.author = "" if anonymized else "[AUTHOR INPUT REQUIRED]"
    doc.core_properties.subject = "Engineering Applications of Artificial Intelligence submission package"
    doc.core_properties.keywords = ", ".join(KEYWORDS)
    return doc


def add_title(doc: Document, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p2.add_run(subtitle)
        r.italic = True; r.font.name = "Arial"; r.font.size = Pt(10); r.font.color.rgb = RGBColor(71, 85, 105)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.55)
    p.paragraph_format.widow_control = True
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None, caption: str | None = None) -> None:
    if caption:
        p = doc.add_paragraph(style="Caption EAAI")
        p.add_run(caption).bold = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for idx, header in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], "17365D")
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color="FFFFFF", size=8)
        if widths:
            table.rows[0].cells[idx].width = Cm(widths[idx])
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for r_index, row_data in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for idx, value in enumerate(row_data):
            if r_index % 2:
                shade_cell(row.cells[idx], "F4F7FA")
            set_cell_text(row.cells[idx], value, size=7.8)
            if widths:
                row.cells[idx].width = Cm(widths[idx])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc: Document, filename: str, number: int, width_inches: float = 6.55) -> None:
    path = optimized_docx_image(FIG_MAIN / f"{filename}.png")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width_inches))
    cap = doc.add_paragraph(style="Caption EAAI")
    cap.add_run(f"Fig. {number}. ").bold = True
    cap.add_run(CAPTIONS[filename])


def add_supp_figure(doc: Document, filename: str, number: str, caption: str, width_inches: float = 6.55) -> None:
    path = optimized_docx_image(FIG_SUPP / f"{filename}.png")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width_inches))
    cap = doc.add_paragraph(style="Caption EAAI")
    cap.add_run(f"Fig. {number}. ").bold = True
    cap.add_run(caption)


def optimized_docx_image(source: Path) -> Path:
    """为旧版 Word 生成轻量内嵌副本；正式独立图件仍保持 600 dpi。"""
    FIG_DOCX.mkdir(parents=True, exist_ok=True)
    target = FIG_DOCX / source.name
    with Image.open(source) as image:
        image = image.convert("RGB")
        max_width = 2200
        if image.width > max_width:
            height = round(image.height * max_width / image.width)
            image = image.resize((max_width, height), Image.Resampling.LANCZOS)
        image.save(target, format="PNG", dpi=(330, 330), optimize=True)
    return target


def format_reference(index: int, ref: dict) -> str:
    authors = ref.get("authors", [])
    author_text = ", ".join(authors[:6]) + (", et al." if len(authors) > 6 else "")
    title = ref.get("title", "")
    venue = ref.get("venue", "")
    year = ref.get("year", "")
    volume = ref.get("volume", "")
    pages = ref.get("article_number_or_pages", "")
    doi = ref.get("doi", "")
    url = ref.get("url", "")
    tail = f" {venue} {volume}".strip()
    if pages:
        tail += f", {pages}"
    tail += f" ({year})."
    if doi:
        tail += f" https://doi.org/{doi}"
    elif url:
        tail += f" {url}"
    return f"[{index}] {author_text}. {title}. {tail}"


def add_references(doc: Document) -> None:
    add_heading(doc, "References", 1)
    for idx, ref in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.7)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(format_reference(idx, ref))


def manuscript_abstract() -> str:
    return (
        "Mountain-road inspection requires an unmanned aerial vehicle to select priority-weighted fixed points while respecting energy, distance, mission-time, dynamic, and depot-return constraints. "
        "We formulate this engineering problem as return-aware combinatorial decision-making and develop a proximal policy optimisation actor–critic with priority-biased self-attention, a pointer decoder, and a deterministic feasibility mask. "
        "The mask removes actions that cannot complete the visit and preserve a feasible return, while the policy handles variable candidate sets within the trained 16-, 20-, and 24-point range. "
        "The frozen evaluation comprises 24 unseen procedural mountain maps (216 tasks) and eight independent Copernicus digital surface model regions (144 tasks), with map-level non-parametric inference. "
        "On unseen procedural maps, the method achieved mean safe priority-weighted coverage of 0.486 versus 0.269 for fixed-slot multilayer-perceptron proximal policy optimisation (Hodges–Lehmann difference 0.217, 95% hierarchical-bootstrap interval 0.196–0.239; Holm-adjusted P=9.54×10−7). "
        "Its coverage was indistinguishable from the attention-matched advantage actor–critic comparator (difference 0.00119; interval −0.00568–0.00689; adjusted P=0.845). "
        "Ant-colony optimisation, simulated annealing, and mixed-integer programming attained higher coverage in some comparisons but required substantially greater online computation. "
        "Removing the return reserve caused the dominant ablation loss (difference 0.371; interval 0.343–0.402). "
        "The results support return-aware neural planning as a fast, reproducible simulation tool, while geographic transfer remains simulation-only and does not constitute flight validation or safety certification."
    )


def build_manuscript() -> Path:
    doc = new_document(TITLE, anonymized=True)
    add_title(doc, TITLE, "Original Research — anonymized manuscript")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Word count (main text, approximate): generated draft; verify after author revision")
    r.font.name = "Arial"; r.font.size = Pt(8); r.font.color.rgb = RGBColor(100, 116, 139)

    add_heading(doc, "Abstract", 1)
    abstract = manuscript_abstract()
    if len(abstract.split()) > 250:
        raise ValueError(f"Abstract is {len(abstract.split())} words")
    p = doc.add_paragraph(abstract); p.paragraph_format.first_line_indent = Cm(0)
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0)
    p.add_run("Keywords: ").bold = True
    p.add_run("; ".join(KEYWORDS))

    add_heading(doc, "1. Introduction", 1)
    add_body(doc, "Mountain-road infrastructure inspection couples route construction with operational safety. Fixed inspection points can differ in priority, elevation, road context, and access cost, while the aircraft must retain sufficient energy, distance, and mission time to return to a depot. These constraints make the task different from point-to-point navigation: the planner must decide which points to service, in which order, and when to stop. Structural-inspection studies have consequently treated coverage and flight length as coupled planning objectives [27], and recent work has formulated periodic power-inspection allocation and routing with endurance-aware action pruning [1].")
    add_body(doc, "Unmanned aerial vehicle navigation research has applied reinforcement learning to collision avoidance, motion control, and path construction [2,3,7]. Engineering Applications of Artificial Intelligence has also reported reinforcement-learning formulations for wind-disturbed quadrotor navigation [2], field-tested closed-loop path planning [3], uneven-terrain risk mapping [4], multi-aircraft energy-aware trajectories [5], and curriculum-guided manoeuvre decisions [6]. These studies show that learning-based planners can encode complex state information, but they do not remove the need to expose feasibility logic, computational cost, and external-validity limits. The general challenges of real-world reinforcement learning—including distribution shift, constraint specification, and reproducible evaluation—remain directly relevant [11].")
    add_body(doc, "Two algorithmic ideas are useful for the present combinatorial setting. Proximal policy optimisation (PPO) stabilises actor updates through a clipped surrogate objective [12], whereas pointer networks define distributions over input positions and therefore avoid a fixed semantic output class for each action [13]. Multi-head attention represents interactions among candidate nodes [14], and attention-based construction policies have solved routing problems with variable candidate sets [17]. However, an attention policy alone does not guarantee that a selected inspection point leaves enough resources for a depot return. Safe-reinforcement-learning literature distinguishes reward penalties from mechanisms that restrict or constrain actions [18,19]. For return-critical inspection, making feasibility explicit before sampling is operationally clearer than asking a reward function to learn every hard boundary.")
    add_body(doc, "This study addresses a multi-resource fixed-point inspection problem on mountainous roads. The artificial-intelligence contribution is a priority-aware PPO–Pointer actor–critic whose legal-action mask accounts for the visit, dynamic, energy, distance, time, and return conditions at every decision. The engineering contribution is a frozen benchmark that connects procedural mountain maps, public Copernicus digital surface model (DSM) regions [20], wind and model perturbations, classical optimisation baselines, and map-level inference. The study tests three questions: whether pointer attention improves safe coverage relative to a genuinely attention-free fixed-slot PPO; whether PPO differs from an attention-matched advantage actor–critic (A2C); and which return, priority, domain-randomisation, and resource-shaping components are supported by ablation evidence.")
    add_body(doc, "The evidence does not support a universal coverage champion claim. Ant-colony optimisation (ACO), simulated annealing (SA), and mixed-integer linear programming (MILP) achieve higher coverage in parts of the benchmark, whereas the neural planners offer shorter online planning times. We therefore report coverage, safety, resource use, computation, training stability, zero-shot geographic simulation transfer, robustness, and ablations as separate evidence dimensions. Real flight, safety certification, causal attribution of PPO mechanisms, and generalisation beyond the trained 16-, 20-, and 24-point problem sizes are outside the supported claim boundary.")

    add_heading(doc, "2. Related work and research gap", 1)
    add_heading(doc, "2.1 Unmanned aerial vehicle inspection and constrained routing", 2)
    add_body(doc, "Inspection path planning has been approached with sampling-based search, evolutionary optimisation, coverage planning, and learned decision policies. Goal-biased PF-RRT* addresses collision-aware trajectory construction in cluttered environments [8], while optimised structural inspection combines viewpoint design and route optimisation [27]. Such methods can provide strong solution quality, but their online computation and solver-status interpretation matter when plans must be regenerated frequently. The periodic fixed-nest inspection formulation in EAAI is the closest published engineering neighbour to the present task [1]; it uses dynamic pruning and topology-aware reinforcement learning for long-horizon power inspection. The present work instead studies a single-mission, priority-weighted, multi-resource mountain-road setting and evaluates a return-aware mask against attention-matched, attention-free, classical, and ablated comparators.")
    add_heading(doc, "2.2 Attention-based reinforcement learning for combinatorial decisions", 2)
    add_body(doc, "Pointer networks map encoded input elements to sequential output positions [13]. When coupled with self-attention [14], the representation can aggregate interactions among spatially distributed candidates. Attention-based routing policies have shown that such construction models can learn reusable heuristics [17]. We use this design only within the trained node-count support and do not infer scale extrapolation. PPO provides the primary policy-update rule [12], generalised advantage estimation reduces variance in return estimates [15], and A2C provides an attention-matched actor–critic comparator rooted in advantage-based asynchronous methods [16]. This pairing isolates the optimiser-level comparison more cleanly than comparing PPO–Pointer only with a fixed multilayer perceptron.")
    add_heading(doc, "2.3 Safety constraints, uncertainty, and statistical comparison", 2)
    add_body(doc, "Safe reinforcement learning covers action restriction, risk-sensitive objectives, constrained optimisation, and verification-oriented approaches [18,19]. Our mask is an engineering feasibility filter rather than a formal constrained-policy-optimisation guarantee: it evaluates the simulated transition and return requirements before an action can be sampled. Domain randomisation provides a practical strategy for exposing a policy to parameter variation during training [10], but it cannot establish real-world transfer on its own. We therefore separate known shifts from hidden model/perception mismatch and label DSM evaluation as geographic simulation transfer.")
    add_body(doc, "Algorithm comparisons across multiple maps require the independent unit to match the scope of generalisation. Rank-based tests avoid distributional assumptions [21,22], while paired map-level comparisons preserve correspondence across algorithms. We use Friedman omnibus tests, two-sided paired Wilcoxon tests, Holm familywise correction [26], Hodges–Lehmann paired location effects [23], and a hierarchical bootstrap over maps [24]. This design follows the multiple-dataset logic described for learning-algorithm comparisons [25] and prevents route or task repetitions from inflating the apparent sample size.")

    add_heading(doc, "3. Mountain-road inspection problem formulation", 1)
    add_heading(doc, "3.1 Task, state, and route", 2)
    add_body(doc, "A task contains a depot and n fixed inspection points. Point i has three-dimensional coordinates xi=(xi,yi,zi), a discrete priority wi, and a visited indicator zi. The aircraft state contains its current position and direction, remaining energy, remaining flight distance, remaining mission time, visited fraction, accumulated priority coverage, and local wind. A route π is a finite sequence of selected inspection points followed by depot return. The trained task support uses n∈{16,20,24}; these are task sizes seen during training.")
    add_body(doc, "The primary confirmatory outcome is safe priority-weighted coverage, SWC(π)=Isafe(π)Σi wi zi/Σi wi. The safety gate Isafe equals one only when the route satisfies the hard feasibility definition and completes a valid depot return; otherwise SWC is zero. Secondary outcomes include unweighted coverage, priority-stratum coverage, safe-route rate, return rate, energy/distance/time utilisation, online planning time, and robustness retention. The post-hoc seven-dimensional 100-point score is restricted to Supplementary Material because its weights and normalisation were selected after the confirmatory outcome.")
    add_heading(doc, "3.2 Multi-resource transition and return condition", 2)
    add_body(doc, "For each candidate point, the simulator computes the outgoing segment and the required return segment under terrain, wind, kinematic, service-time, and reserve assumptions. An action remains legal only if the visit is unserved, the transition respects the dynamic checks, and the predicted energy, distance, and time after the visit can still support the required return. The logical mask is therefore the intersection of visit, dynamics, energy, distance, time, and return-feasibility conditions. Returning to the depot is available when no further inspection action should be taken. In the no-return-reserve ablation, the policy can propose an action that does not preserve the return reserve; the environment records the violation and terminates the stranded route without moving the aircraft.")
    add_heading(doc, "3.3 Aircraft and environmental proxy", 2)
    add_body(doc, "The frozen aircraft proxy uses a 115.2 Wh battery, 13 m s−1 reference cruise speed, 15 m s−1 maximum horizontal speed, 6 m s−1 ascent/descent limits, and a 25% energy reserve. Power proxies are 172.8 W in hover, 138.24 W in cruise, 216 W in climb, and 138.24 W in descent. Each inspection requires 20 s service time. Terrain clearance is 18 m, terrain sampling is 10 m, and a 1.10 safety factor is applied. These values describe the simulation proxy; they are not a certified performance envelope for any physical aircraft. Rotorcraft propulsion depends nonlinearly on speed and flight regime [9], so the simplified proxy is treated as a reproducible test model rather than a high-fidelity powertrain digital twin.")

    add_heading(doc, "4. Return-aware proximal policy optimisation with pointer attention", 1)
    add_figure(doc, "Fig01_method", 1)
    add_heading(doc, "4.1 Priority-aware node and aircraft representation", 2)
    add_body(doc, "Each node is represented by 15 features: relative three-dimensional position, normalised priority, visited and depot flags, outgoing distance/energy/time utilisation, return distance/energy/time utilisation, and outgoing wind components. The 14-dimensional aircraft vector contains relative current position, previous direction, remaining energy/distance/time, visited fraction, priority coverage, and current wind. Continuous resource quantities are normalised using task budgets. Candidate padding is used only to batch variable tasks; invalid padding never becomes a legal action.")
    add_body(doc, "A linear projection maps node features to a 128-dimensional embedding. The encoder applies four-head self-attention with an additive priority bias scaled by 0.5, followed by a feed-forward block and layer normalisation. This bias changes attention scores but does not itself enforce visiting a high-priority point. The no-priority-bias ablation removes the additive term while keeping the remaining architecture and training protocol fixed.")
    add_heading(doc, "4.2 Pointer actor and critic", 2)
    add_body(doc, "The actor maps the aircraft vector to a query, applies a masked multi-head glimpse over encoded nodes, and gates the raw and context-enriched queries. Pointer logits are produced by a tanh-scaled compatibility function with scale 10. The legal mask is applied at the glimpse and final-logit stages. Sampling or greedy decoding therefore operates only over actions declared feasible by the simulator. The critic concatenates the query with masked means over remaining and legal nodes and predicts a scalar value through a multilayer perceptron.")
    add_body(doc, "The audited full seed-42 checkpoint contains 341,505 trainable parameters. The attention-free traditional comparator uses a fixed 24-point-plus-depot slot representation: 25×15 node inputs plus the 14 aircraft inputs form a 389-dimensional actor input, followed by 256-unit hidden layers and 25 action logits. Its critic additionally receives the validity and legality masks. This Flat-MLP PPO comparator contains no attention operation; the previously explored ppo_mlp variant is excluded because it retained attention and therefore did not satisfy the intended baseline identity.")
    add_heading(doc, "4.3 PPO training and validation selection", 2)
    add_body(doc, "PPO uses a learning rate of 1×10−4, clipping parameter 0.2, discount factor 0.99, generalised-advantage parameter 0.95, value coefficient 0.5, gradient-norm limit 1.0, five optimisation epochs per update, 16 episodes per update, and minibatches of 128. Entropy regularisation decreases from 0.02 to 0.002, and the target Kullback–Leibler divergence is 0.02. The primary reward is priority gain with a small secondary unweighted-coverage term. Resource shaping uses a small priority-scaled mean incremental utilisation term; hard feasibility remains in the action mask rather than the reward.")
    add_body(doc, "Training domain randomisation samples initial state of charge from 0.80–1.00, resource-scale factors from 0.85–1.00, wind scale from 0.80–1.20, coordinate rotation within ±15°, and vertical-wind bias within ±1 m s−1. For each learning variant, five seeds (42–46) are trained for 3,000 episodes. The best safe checkpoint is selected on 108 validation tasks. Across seven learning variants, the paper model set comprises 35 checkpoints and 105,000 training episodes. [AUTHOR INPUT REQUIRED: exact operating system, processor/GPU model, PyTorch version, and dependency versions used for the frozen training and evaluation.] ")

    add_heading(doc, "5. Experimental and statistical protocol", 1)
    protocol_rows = [
        ["Training", "72 procedural maps", "648", "16/20/24", "Five seeds; domain randomisation"],
        ["Validation", "12 procedural maps", "108", "16/20/24", "Checkpoint selection only"],
        ["Unseen synthetic test", "24 procedural maps", "216", "16/20/24", "Confirmatory map-level inference"],
        ["Geographic DSM test", "8 independent regions; two road contexts each", "144", "16/20/24", "Zero-shot simulation transfer"],
    ]
    add_table(doc, ["Split", "Map basis", "Tasks", "Point counts", "Role"], protocol_rows, [3.0, 5.0, 2.0, 2.5, 5.0], "Table 1. Frozen map and task protocol.")
    add_heading(doc, "5.1 Procedural mountain maps and tasks", 2)
    add_body(doc, "Procedural terrain uses 267×267 grids with 30 m cells, six fractal octaves, Hurst parameters 0.55–0.85, and multiple road topologies. Tasks combine moderate, hard, and extreme radius bands with energy, distance, time, or mixed binding constraints. Candidate points are sampled from road contexts with minimum 180 m depot distance, 120 m node spacing, and 90 m road sampling. The train, validation, and unseen-test map seeds are disjoint. All evaluated node counts occur in training; consequently, the test estimates unseen-map generalisation, not scale extrapolation.")
    add_heading(doc, "5.2 Geographic digital surface model simulations", 2)
    add_body(doc, "The geographic set contains four Chinese and four global mountain regions: Qinling, Taihang, Hengduan, Guizhou karst, the Alps, Colorado Rockies, Peruvian Andes, and New Zealand Southern Alps. Each approximately 8 km crop uses the public Copernicus DEM GLO-30 product [20], requires at least 300 m relief, at least 6 km of road, and at least three road branches, and contributes two road contexts. Copernicus DEM is a digital surface model; vegetation and built objects may contribute to elevation. No geographic region is used for policy training or validation. The evaluation is therefore described as zero-shot geographic DSM simulation transfer.")
    add_heading(doc, "5.3 Comparators and ablations", 2)
    add_body(doc, "The principal learning comparators are attention-matched A2C+Pointer and attention-free Flat-MLP PPO. Deterministic heuristics include nearest-feasible and priority-resource greedy planning. Stochastic classical comparators include ACO, genetic algorithm (GA), and SA; MILP provides an optimisation reference under the frozen formulation. Supplementary comparisons include A*, particle swarm optimisation, and an exact Pareto dynamic programme on the tractable subset. Four ablations remove priority bias, domain randomisation, resource shaping, or the return reserve. Historical ppo_mlp results are excluded from every paper table and figure.")
    add_heading(doc, "5.4 Robustness design", 2)
    add_body(doc, "Robustness has two layers. Known shifts alter variables represented by the training randomisation envelope, including wind and power scaling. Hidden mismatch introduces perturbations not explicitly represented during training, including DEM error, localisation error, power-model mismatch, and wind mismatch. Performance retention is computed relative to each model's own reference condition; it should not be read as absolute cross-model performance. The safe-route share is reported alongside retention.")
    add_heading(doc, "5.5 Statistical analysis", 2)
    add_body(doc, "The map is the independent unit. For each metric and algorithm, task, road, seed, and route repetitions are first aggregated within maps according to the frozen analysis definition. A Friedman test evaluates each prespecified algorithm family [21,25]. When the omnibus comparison is retained, PPO+Pointer is compared with each prespecified comparator using two-sided paired Wilcoxon signed-rank tests [22]. Holm adjustment controls the familywise error rate within each family [26]. We report Hodges–Lehmann paired location effects [23], rank-biserial direction, and 95% intervals from 10,000 hierarchical-bootstrap replicates whose outer resampling unit is the map [24]. Exact P values are reported where representable. Statistical significance is not used to replace effect magnitude or engineering interpretation.")
    add_body(doc, "The confirmatory outcome is safe priority-weighted coverage. The seven-dimensional operational score and its weight sensitivity are post-hoc summaries and remain in Supplementary Material. No task, road, seed, or route count is presented as the independent sample size for map-level tests. All analyses use the locked 21,648-record result file and the frozen v3.2.14 scripts; no training, evaluation, or statistical procedure was rerun for manuscript preparation.")

    add_heading(doc, "6. Results", 1)
    add_heading(doc, "6.1 Safe priority-weighted coverage", 2)
    add_figure(doc, "Fig02_coverage_priority", 2)
    full_syn = find_row(DESCRIPTIVE, algorithm="full", result_family="synthetic_learning", metric="safe_weighted_coverage")
    trad_syn = find_row(DESCRIPTIVE, algorithm="traditional_ppo", result_family="synthetic_learning", metric="safe_weighted_coverage")
    a2c_syn = find_row(DESCRIPTIVE, algorithm="a2c_pointer", result_family="synthetic_learning", metric="safe_weighted_coverage")
    syn_om = find_row(OMNIBUS, statistical_family="synthetic_main_algorithms")
    p_trad = find_row(PAIRWISE, statistical_family="synthetic_main_algorithms", comparator="traditional_ppo")
    p_a2c = find_row(PAIRWISE, statistical_family="synthetic_main_algorithms", comparator="a2c_pointer")
    add_body(doc, f"The main-algorithm distributions differed across 24 unseen synthetic maps (Friedman statistic {f3(syn_om['statistic'])}, P={sci(syn_om['p_value'])}; Fig. 2a). PPO+Pointer achieved mean safe priority-weighted coverage {f3(full_syn['mean'])}, compared with {f3(trad_syn['mean'])} for Flat-MLP PPO. The paired Hodges–Lehmann difference was {f3(p_trad['hodges_lehmann'])} (95% hierarchical-bootstrap interval {f3(p_trad['bootstrap_ci_low'])}–{f3(p_trad['bootstrap_ci_high'])}; Holm-adjusted P={sci(p_trad['p_holm'])}). Thus, replacing the fixed-slot multilayer perceptron with the pointer-attention architecture was associated with a large improvement under the frozen protocol.")
    add_body(doc, f"PPO+Pointer and A2C+Pointer had nearly identical mean coverage ({f3(full_syn['mean'])} and {f3(a2c_syn['mean'])}, respectively). Their paired difference was {f3(p_a2c['hodges_lehmann'])} (interval {f3(p_a2c['bootstrap_ci_low'])} to {f3(p_a2c['bootstrap_ci_high'])}; adjusted P={f3(p_a2c['p_holm'])}). This non-significant contrast does not establish formal equivalence; it shows that the confirmatory outcome did not distinguish the two attention-matched actor–critic methods.")
    add_body(doc, "Classical optimisation exposed a different trade-off. Mean unseen-map coverage was 0.553 for ACO, 0.544 for SA, and 0.571 for MILP, each above PPO+Pointer. The full-minus-comparator Hodges–Lehmann effects were −0.066 for ACO, −0.056 for SA, and −0.084 for MILP (all Holm-adjusted P=9.54×10−7). GA was close to PPO+Pointer (mean 0.487; adjusted P=0.922). Consequently, coverage alone does not support a PPO champion conclusion. Priority-stratum patterns were similar with and without priority bias (Fig. 2b), consistent with the non-significant aggregate priority-bias ablation reported below.")
    add_heading(doc, "6.2 Safety, depot return, and resource utilisation", 2)
    add_figure(doc, "Fig03_safety_resources", 3)
    add_body(doc, "All PPO+Pointer routes in the primary unseen synthetic and geographic DSM evaluations satisfied the frozen safe-route and depot-return definitions. Under known shifts, the safety and return differences relative to A2C+Pointer and Flat-MLP PPO were centred at zero because these comparators also retained full primary feasibility. Hidden mismatch reduced absolute safe and return rates, but the mean percentage-point effects favoured PPO+Pointer over the two principal learning comparators (Fig. 3a). These observations describe the simulated feasibility gate; they are not evidence of real-aircraft safety.")
    add_body(doc, "Among safe routes, median energy, distance, and mission-time utilisation remained below one for all three learning models (Fig. 3b). Flat-MLP PPO used a similar or larger fraction of some budgets while visiting substantially less priority weight, which indicates inefficient route selection rather than a simple resource shortage. The data do not identify a causal contribution for each energy, distance, time, or dynamics submask because those submasks were not independently ablated.")
    add_heading(doc, "6.3 Online planning time and quality–time trade-off", 2)
    add_figure(doc, "Fig04_time_tradeoff", 4)
    add_body(doc, "On unseen synthetic tasks, the 95th-percentile online planning time was 2.07 s for PPO+Pointer, 1.96 s for A2C+Pointer, and 1.11 s for Flat-MLP PPO. The corresponding geographic DSM values were 2.13, 2.09, and 1.23 s. The attention-free network was faster but produced substantially lower safe priority-weighted coverage. PPO and A2C had comparable latency and coverage.")
    add_body(doc, "The classical methods occupied a different computational range. ACO required 134.29 s at the synthetic 95th percentile and 82.58 s on geographic DSM tasks; MILP required 46.22 and 39.95 s, respectively. Both attained higher coverage dimensions than the learned policies. These values are hardware- and implementation-specific, but the order-of-magnitude separation supports an explicit online quality–time trade-off rather than a universal ranking.")
    add_heading(doc, "6.4 Training stability and sample efficiency", 2)
    add_figure(doc, "Fig05_training", 5)
    add_body(doc, "Across five seeds, PPO+Pointer and A2C+Pointer reached similar final safe weighted coverage, whereas Flat-MLP PPO converged to a substantially lower level (Fig. 5a). The frozen D6 training-stability score was 0.9860 for PPO+Pointer and 0.9701 for A2C+Pointer. The D7 sample-efficiency scores were 0.9654 and 0.9466, respectively (Fig. 5b). These descriptive training dimensions favour PPO+Pointer, but they do not prove which PPO mechanism caused the difference because the optimiser, trajectories, and value updates were not independently randomised.")
    add_heading(doc, "6.5 Unseen procedural maps and geographic DSM simulations", 2)
    add_figure(doc, "Fig06_transfer", 6)
    real_full = find_row(DESCRIPTIVE, algorithm="full", result_family="real_learning", metric="safe_weighted_coverage")
    real_a2c = find_row(DESCRIPTIVE, algorithm="a2c_pointer", result_family="real_learning", metric="safe_weighted_coverage")
    real_trad = find_row(DESCRIPTIVE, algorithm="traditional_ppo", result_family="real_learning", metric="safe_weighted_coverage")
    real_aco = find_row(DESCRIPTIVE, algorithm="aco", result_family="real_baselines", metric="safe_weighted_coverage")
    real_milp = find_row(DESCRIPTIVE, algorithm="milp", result_family="real_baselines", metric="safe_weighted_coverage")
    add_body(doc, f"Without geographic fine-tuning, PPO+Pointer achieved mean safe weighted coverage {f3(real_full['mean'])} across eight independent DSM regions. A2C+Pointer was nearly identical ({f3(real_a2c['mean'])}), while Flat-MLP PPO reached {f3(real_trad['mean'])}. ACO and MILP again produced higher coverage ({f3(real_aco['mean'])} and {f3(real_milp['mean'])}, respectively). The paired PPO–A2C effect was 0.000496 (interval −0.00904 to 0.00896; adjusted P=0.742), whereas PPO exceeded Flat-MLP PPO by 0.240 (interval 0.212–0.262; adjusted P=0.0469). The representative Taihang panel visualises planned routes over DSM elevation (Fig. 6b).")
    add_body(doc, "These results show transfer across map seeds and geographic elevation/road contexts within the simulator. They do not show transfer from simulation to physical flight, do not validate perception or control hardware, and do not establish performance on point counts outside the training support.")
    add_heading(doc, "6.6 Known shifts and hidden mismatch", 2)
    add_figure(doc, "Fig07_robustness", 7)
    add_body(doc, "Mean safe weighted coverage under known shifts was 0.479 for PPO+Pointer and 0.480 for A2C+Pointer. Under hidden model/perception mismatch, the means decreased to 0.429 and 0.418, respectively; Flat-MLP PPO reached 0.231. Retention patterns were generally high when normalised to each model's own reference, but safe-route shares revealed larger condition-dependent degradation (Fig. 7). The no-domain-randomisation variant achieved 0.428 under hidden mismatch, close to the full model, so this finite perturbation catalogue does not isolate an independent benefit of domain randomisation.")
    add_heading(doc, "6.7 Component ablations", 2)
    add_figure(doc, "Fig08_ablation", 8)
    syn_ab_om = find_row(OMNIBUS, statistical_family="synthetic_ablations")
    ab_ret = find_row(PAIRWISE, statistical_family="synthetic_ablations", comparator="no_return_reserve")
    add_body(doc, f"The synthetic ablation family differed across 24 maps (Friedman statistic {f3(syn_ab_om['statistic'])}, P={sci(syn_ab_om['p_value'])}). Removing the return reserve reduced safe weighted coverage from 0.486 to 0.113. The Hodges–Lehmann full-minus-ablation effect was {f3(ab_ret['hodges_lehmann'])} (interval {f3(ab_ret['bootstrap_ci_low'])}–{f3(ab_ret['bootstrap_ci_high'])}; adjusted P={sci(ab_ret['p_holm'])}). A similarly large effect occurred in the geographic DSM set (Fig. 8). This is the clearest component-level evidence in the study.")
    add_body(doc, "Removing priority bias, domain randomisation, or resource shaping produced synthetic effects of 0.00133, 0.00122, and 0.00181, respectively; none remained significant after Holm correction. These findings do not prove that the components are universally unnecessary. They show that their independent removal was not distinguished by the confirmatory coverage outcome under this training budget and evaluation set. Because energy, distance, time, and dynamics submasks were not individually removed, no separate causal contribution is assigned to those mask components.")

    add_heading(doc, "7. Discussion and limitations", 1)
    add_heading(doc, "7.1 Principal findings", 2)
    add_body(doc, "Three observations are supported. First, the return-aware pointer architecture substantially outperformed an attention-free fixed-slot PPO on safe priority-weighted coverage in both unseen procedural and geographic DSM simulations. Second, PPO+Pointer and attention-matched A2C+Pointer were not distinguishable on the confirmatory coverage outcome, although PPO+Pointer had higher frozen training-stability and sample-efficiency scores. Third, the return reserve was the only individually tested component whose removal caused a large, consistent loss. Classical optimisers achieved higher coverage in several comparisons but required far more online planning time.")
    add_heading(doc, "7.2 Interpretation of the architecture and optimiser evidence", 2)
    add_body(doc, "The comparison with Flat-MLP PPO is consistent with the interpretation that a variable-candidate pointer representation is better suited to structured inspection decisions than a fixed output slot mapping. This interpretation is supported by architecture and performance differences, but it is not a causal decomposition: the models differ in representation, parameter sharing, and candidate interaction. Pointer networks and attention-based routing provide a plausible mechanism for representing candidate relationships [13,17], yet the present experiments do not directly measure learned relational features.")
    add_body(doc, "The PPO–A2C comparison is more controlled because both use the pointer architecture. Their similar confirmatory coverage indicates that attention and feasibility structure dominate final primary-outcome differences in this benchmark. PPO's stronger D6/D7 profile suggests more stable or efficient learning under the selected hyperparameters, consistent with the intent of clipped updates [12]. However, no optimiser-by-hyperparameter factorial experiment was performed; the result should therefore be read as protocol-specific performance, not a general proof that PPO is intrinsically superior to A2C.")
    add_heading(doc, "7.3 Engineering use conditions and classical optimisation", 2)
    add_body(doc, "The neural policies are most defensible when repeated, low-latency replanning is required and a small coverage concession relative to slower solvers is acceptable. MILP, ACO, and SA remain attractive when solution quality dominates latency or when solver-status information is operationally valuable. The study does not treat these baselines as obsolete; instead, they anchor the attainable coverage range and reveal where the learned policy trades quality for time. A hybrid workflow could use a neural policy for immediate plans and a slower optimiser for offline auditing or warm-start refinement, but this is a proposed application pattern rather than a tested result.")
    add_heading(doc, "7.4 Return-aware feasibility and safety boundaries", 2)
    add_body(doc, "The return-reserve ablation supports the value of checking post-visit return feasibility before action selection. This is an action-restriction strategy within safe reinforcement learning [19], not a formal guarantee about all uncertainty sources. The simulated mask relies on the correctness of the terrain, wind, power, timing, and state estimates. Hidden mismatch reduced safe-route shares, demonstrating that nominal feasibility can degrade when these estimates are wrong. Formal constrained-policy methods [18], online uncertainty estimation, reachability analysis, and flight-controller verification could complement the present planner, but none was evaluated here.")
    add_heading(doc, "7.5 Transfer, reproducibility, and limitations", 2)
    add_body(doc, "The independent Copernicus regions broaden the terrain and road contexts beyond procedural map generation, and the locked package makes the transfer claim auditable. Nevertheless, Copernicus GLO-30 is a DSM with 30 m nominal sampling [20]; it cannot resolve all local obstacles, wires, vegetation motion, or road-edge hazards relevant to flight. The experiments also use synthetic wind and simplified power proxies. Accordingly, geographic DSM transfer is a simulation-domain result, not sim-to-real validation.")
    add_body(doc, "Five additional limitations constrain interpretation. First, task sizes are limited to node counts observed during training. Second, only eight geographic regions are independent real-terrain units, which limits precision and diversity. Third, the policy was not deployed on an aircraft, and no perception, control, communication, or regulatory subsystem was tested. Fourth, resource models use transparent proxies rather than vehicle-specific identification. Fifth, the ablation design does not isolate every mask subcondition or interaction. Future work should preregister confirmatory metrics, add out-of-range scale tests, evaluate calibrated uncertainty and stronger constrained-learning baselines, and conduct staged hardware-in-the-loop and flight validation with independent safety review.")

    add_heading(doc, "8. Conclusions", 1)
    add_body(doc, "A return-aware PPO–Pointer planner was evaluated for priority-weighted, multi-resource mountain-road inspection. The method strongly improved safe weighted coverage over an attention-free fixed-slot PPO and preserved short online planning times. Its confirmatory coverage was close to that of A2C+Pointer, while the frozen training dimensions favoured PPO. ACO, SA, and MILP remained stronger on parts of the coverage outcome at materially greater computation. Removing the return reserve produced the dominant ablation loss, making explicit return feasibility the clearest supported safety-related contribution. The findings justify reproducible simulation use and further engineering validation; they do not establish real-flight performance, formal safety certification, out-of-range scale generalisation, or causal superiority of PPO mechanisms.")

    add_heading(doc, "Data and code availability", 1)
    add_body(doc, "The anonymised reproducibility package accompanying this submission contains the frozen protocol, task/evaluation manifests, result tables, statistical Source Data, figure Source Data, reconstruction scripts, and file hashes. Raw Copernicus assets are not redistributed unless their current licence permits redistribution; the package provides source links, region identifiers, and reconstruction instructions. [AUTHOR INPUT REQUIRED: anonymous repository URL or DOI for peer review.] After acceptance, [AUTHOR INPUT REQUIRED: permanent public repository and licence].")
    add_heading(doc, "Declaration of generative artificial intelligence use", 1)
    add_body(doc, "During preparation of this work, the authors used OpenAI Codex to assist with language drafting, document assembly, and English relabelling of plots from frozen Source Data. No generative image system was used for scientific figures, and the tool did not generate, modify, or analyse experimental results. The authors reviewed and edited all outputs and remain responsible for the manuscript, evidence, citations, and conclusions. [AUTHOR CHECK: verify the target journal's current required wording and submission-field placement immediately before submission.]")
    add_heading(doc, "Funding", 1)
    add_body(doc, "[AUTHOR INPUT REQUIRED: funding bodies, grant numbers, and recipient initials; state ‘no specific funding’ only if accurate.]")
    add_heading(doc, "CRediT authorship contribution statement", 1)
    add_body(doc, "[AUTHOR INPUT REQUIRED: author names and verified CRediT roles.]")
    add_heading(doc, "Declaration of competing interest", 1)
    add_body(doc, "[AUTHOR INPUT REQUIRED: verified competing-interest statement.]")
    add_heading(doc, "Acknowledgements", 1)
    add_body(doc, "[AUTHOR INPUT REQUIRED: acknowledgements, or confirm that this section should be omitted.]")

    add_references(doc)
    add_heading(doc, "Author Verification Queue", 1)
    add_bullets(doc, [
        "[AUTHOR CHECK 01] Provide author names, affiliations, ORCID identifiers, and corresponding-author details in the separate title page.",
        "[AUTHOR CHECK 02] Confirm exact software, dependency, operating-system, processor, and GPU versions for the frozen runs.",
        "[AUTHOR CHECK 03] Provide the anonymous repository link for review and choose the post-acceptance permanent repository and licence.",
        "[AUTHOR CHECK 04] Confirm Copernicus derivative-data redistribution terms; otherwise retain links and reconstruction instructions only.",
        "[AUTHOR CHECK 05] Supply verified funding, CRediT roles, competing-interest statement, and acknowledgements.",
        "[AUTHOR CHECK 06] Recheck EAAI's current generative-AI disclosure wording and submission-field placement at submission.",
    ])
    path = DOCS / "EAAI_manuscript_anonymized.docx"
    doc.save(path)
    return path


def build_title_page() -> Path:
    doc = new_document(TITLE)
    add_title(doc, TITLE, "Original Research — title page")
    add_heading(doc, "Authors and affiliations", 1)
    add_body(doc, "[AUTHOR INPUT REQUIRED: full author names in publication order, academic degrees if required, and superscript affiliation mapping.]")
    add_body(doc, "[AUTHOR INPUT REQUIRED: complete institutional affiliations, city, postal code, and country.]")
    add_heading(doc, "Corresponding author", 1)
    add_body(doc, "[AUTHOR INPUT REQUIRED: name, institutional postal address, email, and telephone if requested by the submission system.]")
    add_heading(doc, "ORCID", 1)
    add_body(doc, "[AUTHOR INPUT REQUIRED: verified ORCID identifiers.]")
    add_heading(doc, "Manuscript metadata", 1)
    add_table(doc, ["Field", "Value"], [
        ["Target journal", "Engineering Applications of Artificial Intelligence"],
        ["Article type", "Original Research"],
        ["Short title", SHORT_TITLE],
        ["Keywords", "; ".join(KEYWORDS)],
        ["Figures", "8 main figures; 8 supplementary figures"],
        ["Tables", "1 principal protocol table plus supplementary audit tables"],
    ], [5.0, 11.0])
    add_heading(doc, "Declarations", 1)
    add_body(doc, "Funding: [AUTHOR INPUT REQUIRED: verified statement.]")
    add_body(doc, "Competing interests: [AUTHOR INPUT REQUIRED: verified statement.]")
    add_body(doc, "Author contributions: [AUTHOR INPUT REQUIRED: verified CRediT statement.]")
    add_body(doc, "Data/code repository: [AUTHOR INPUT REQUIRED: anonymous link for review; permanent link after acceptance.]")
    path = DOCS / "EAAI_title_page.docx"; doc.save(path); return path


def build_highlights() -> Path:
    doc = new_document("EAAI Highlights")
    add_title(doc, "Highlights", "Engineering Applications of Artificial Intelligence")
    add_bullets(doc, HIGHLIGHTS)
    add_heading(doc, "Compliance note", 1)
    add_body(doc, "Five bullet points are supplied. Each contains no more than 85 characters, including spaces.")
    path = DOCS / "EAAI_highlights.docx"; doc.save(path); return path


def build_cover_letter() -> Path:
    doc = new_document("EAAI Cover Letter")
    add_title(doc, "Cover Letter")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("[AUTHOR INPUT REQUIRED: submission date]")
    add_body(doc, "Dear Editor-in-Chief,")
    add_body(doc, f"We submit the Original Research manuscript entitled “{TITLE}” for consideration in Engineering Applications of Artificial Intelligence.")
    add_body(doc, "The engineering problem is online planning for fixed, priority-weighted mountain-road inspection points under simultaneous energy, distance, mission-time, dynamics, and depot-return constraints. The artificial-intelligence contribution is a priority-aware PPO–Pointer actor–critic with a deterministic return-aware legal-action mask. The engineering contribution is a locked, reproducible evaluation spanning unseen procedural mountain maps, independent public Copernicus DSM regions, classical optimisers, robustness layers, and component ablations.")
    add_body(doc, "The primary evidence is deliberately bounded. PPO+Pointer substantially exceeds an attention-free fixed-slot PPO, but it does not exceed attention-matched A2C+Pointer on confirmatory safe weighted coverage. ACO, SA, and MILP achieve higher coverage in several comparisons at much higher online computation. The strongest component result is the large loss caused by removing the return reserve. We make no claim of real-flight validation, safety certification, out-of-range scale generalisation, or universal PPO superiority.")
    add_body(doc, "The manuscript is formatted as a single-column, double-anonymised Original Research submission. The abstract is below 250 words, five keywords and five Highlights are supplied, and the package includes data/code availability, funding, competing-interest, authorship, and generative-AI declarations or explicit author placeholders. Scientific figures were redrawn only from frozen Source Data; no generative image tool was used.")
    add_body(doc, "[AUTHOR INPUT REQUIRED: confirm that the work is original, is not under consideration elsewhere, and has been approved by all authors.]")
    add_body(doc, "[AUTHOR INPUT REQUIRED: identify any related manuscripts, preprints, or prior conference versions and explain overlap.]")
    add_body(doc, "Suggested reviewers: [AUTHOR INPUT REQUIRED: names, institutions, institutional emails, expertise, and conflict checks, if the journal requests suggestions.]")
    add_body(doc, "Thank you for considering this manuscript.")
    add_body(doc, "Sincerely,")
    add_body(doc, "[AUTHOR INPUT REQUIRED: corresponding author name and contact information]")
    path = DOCS / "EAAI_cover_letter.docx"; doc.save(path); return path


def build_supplement() -> Path:
    doc = new_document("EAAI Supplementary Material", anonymized=True)
    add_title(doc, "Supplementary Material", TITLE)
    add_heading(doc, "S1. Frozen protocol and reproducibility scope", 1)
    add_body(doc, "This supplement reports secondary diagnostics, additional baseline views, failure-mode summaries, all learning curves, and the post-hoc operational composite. It does not change the confirmatory safe priority-weighted coverage outcome, map-level independent unit, or v3.2.14 protocol. The source result file contains 21,648 records, the final audit passed, and ppo_mlp is absent.")
    add_table(doc, ["Item", "Frozen value", "Claim boundary"], [
        ["Learning checkpoints", "35 paper models", "Seven variants × five seeds"],
        ["Paper training episodes", "105,000", "Historical excluded work not used"],
        ["Unseen procedural maps", "24", "Independent map-level test units"],
        ["Geographic DSM regions", "8", "Simulation transfer only"],
        ["Final route records", "21,648", "No manuscript rerun"],
        ["Hierarchical bootstrap", "10,000 replicates", "Outer unit is map"],
    ], [5.0, 5.0, 7.0], "Table S1. Frozen audit facts.")
    add_heading(doc, "S2. Performance profiles and solver diagnostics", 1)
    add_supp_figure(doc, "FigS01_performance_profile", "S1", "Performance profiles based on regret from the safe priority-weighted coverage leader for each task.")
    add_supp_figure(doc, "FigS03_oracle_compute", "S2", "Oracle regret and online computation. Horizontal intervals show the frozen regret range; planning time is logarithmic.")
    add_heading(doc, "S3. Scenario-stratified and failure-mode results", 1)
    add_supp_figure(doc, "FigS04_scenario_heatmap", "S3", "Scenario-stratified safe priority-weighted coverage across node count, terrain, wind, and resource factors.")
    add_supp_figure(doc, "FigS05_failure_modes", "S4", "Safety and return outcomes across robustness families and perturbation conditions.")
    add_heading(doc, "S4. All-model training evidence", 1)
    add_supp_figure(doc, "FigS06_all_training", "S5", "Training trajectories for all seven learning variants under the frozen 3,000-episode protocol.")
    add_heading(doc, "S5. Post-hoc operational summary and sensitivity", 1)
    add_body(doc, "The seven-dimensional operational score combines D1–D7 after the confirmatory analyses. The arithmetic scores are 76.66 for PPO+Pointer, 74.55 for A2C+Pointer, and 54.47 for Flat-MLP PPO. A 10,000-replicate hierarchical bootstrap gives a mean PPO-minus-A2C difference of 1.99 points (95% interval 0.78–3.21; probability positive 1.00). This score is descriptive: its weights, operational floor, and normalisation do not replace safe priority-weighted coverage, and no individual D4, D6, or D7 contrast remained significant after Holm adjustment in its post-hoc family.")
    add_supp_figure(doc, "FigS07_dimensions", "S6", "Seven-dimensional operational evidence profile. The 100-point aggregate is post-hoc and not a confirmatory champion metric.")
    add_supp_figure(doc, "FigS08_sensitivity", "S7", "Sensitivity of PPO+Pointer ranking to the operational floor and combined D6+D7 training weight.")
    add_heading(doc, "S6. Representative unseen synthetic route", 1)
    add_supp_figure(doc, "FigS09_synthetic_route", "S8", "Representative fixed unseen synthetic task with road context, inspection priorities, and comparator routes.")
    add_heading(doc, "S7. Additional claim boundaries", 1)
    add_bullets(doc, [
        "The geographic evidence is DSM-based simulation, not flight validation.",
        "All node counts were observed during training; scale extrapolation is untested.",
        "The ablations do not isolate energy, distance, time, and dynamics submasks separately.",
        "The aircraft and wind models are engineering proxies, not certified digital twins.",
        "Classical solver status and runtime must be interpreted jointly with route quality.",
        "The 100-point score is a post-hoc summary and is not used in the abstract or primary conclusion.",
    ])
    path = DOCS / "EAAI_supplementary_material.docx"; doc.save(path); return path


def build_traceability() -> Path:
    doc = new_document("Journal Style Profile, Evidence Matrix, and Compliance Audit")
    add_title(doc, "Journal Style Profile, Evidence Matrix, and Compliance Audit", "EAAI Original Research — v3.2.14 evidence lock")
    add_heading(doc, "1. Journal Style Profile", 1)
    add_table(doc, ["Rule", "Level", "Source/evidence", "Confidence", "Applied"], [
        ["Single-column Word submission", "H", "Official EAAI guide", "High", "Yes"],
        ["Double-anonymised manuscript and separate title page", "H", "Official EAAI guide", "High", "Yes"],
        ["Abstract ≤250 words; 1–6 keywords", "H", "Official EAAI guide", "High", "Yes"],
        ["Highlights: 3–5 bullets, each ≤85 characters", "H", "Official EAAI guide", "High", "Yes"],
        ["Explicit AI contribution and engineering application contribution", "H", "EAAI aims/scope and desk-reject rules", "High", "Yes"],
        ["Methods and Results separated; independent Conclusions", "S", "Recent Original Research exemplars", "High", "Yes"],
        ["Problem formulation before learning architecture", "S", "Routing/navigation exemplar pool", "High", "Yes"],
        ["Results ordered from primary quality to safety, time, training, transfer, robustness, ablation", "S", "Stable evidence-first convention adapted to study logic", "Medium", "Yes"],
        ["Graphical abstract", "O", "Encouraged but not required", "High", "Not supplied"],
        ["Post-hoc composite in main claim", "O", "Study-specific", "Low", "No; supplement only"],
    ], [5.2, 1.4, 6.0, 2.0, 2.2])
    add_body(doc, "Macrostructure: Introduction → Related work and gap → Problem formulation → Method → Experimental/statistical protocol → Results → Discussion and limitations → Conclusions. Results paragraphs follow purpose → figure → observation → quantitative evidence → statistics → bounded interpretation. Discussion is organised by findings and claim boundaries rather than citation count.")
    add_body(doc, "Language profile: precise active/passive mixture; past tense for completed experiments; present tense for figure content and established concepts; restrained hedging; no unverified first/novel/unprecedented claim; no universal PPO superiority. Sample papers influenced only section functions, information order, evidence density, and caption roles.")

    add_heading(doc, "2. Fact Map", 1)
    add_table(doc, ["Fact ID", "Statement", "Source", "Type", "Status"], [list(x) for x in FACT_MAP], [1.5, 7.4, 4.2, 2.0, 1.8])
    add_heading(doc, "3. Claim Map and Evidence Matrix", 1)
    add_table(doc, ["Claim", "Maximum statement", "Evidence", "Evidence type", "Max strength", "Boundary"], [list(x) for x in CLAIM_MAP], [1.2, 5.3, 3.6, 2.2, 1.8, 4.2])
    add_heading(doc, "4. Figure Storyline", 1)
    add_table(doc, ["Figure", "Question", "Maximum claim", "Cannot prove"], [
        ["Fig. 1", "How are variable candidates and return constraints coupled?", "Feasibility is filtered before action selection.", "Formal or physical safety"],
        ["Fig. 2", "Does the method improve safe coverage?", "Large gain over Flat-MLP PPO; similarity to A2C.", "Universal coverage dominance"],
        ["Fig. 3", "Are routes safe and resource-aware?", "Primary simulated feasibility and bounded utilisation.", "Certified aircraft margins"],
        ["Fig. 4", "What is the online quality–time trade-off?", "Neural methods are much faster than ACO/MILP here.", "Hardware-independent latency"],
        ["Fig. 5", "How do learning trajectories differ?", "PPO has stronger frozen D6/D7 profile.", "Causal PPO mechanism"],
        ["Fig. 6", "Does the policy transfer across geographic DSM contexts?", "Zero-shot simulation transfer.", "Sim-to-real or flight transfer"],
        ["Fig. 7", "How does performance change under shifts?", "Finite known/hidden perturbation robustness.", "Open-world robustness"],
        ["Fig. 8", "Which components have independent evidence?", "Return reserve is dominant.", "Independent submask effects"],
    ], [1.5, 5.5, 6.2, 4.0])

    add_heading(doc, "5. Final Compliance Audit", 1)
    compliance_rows = [
        ["Scientific", "final_results.jsonl = 21,648; passed=true; ppo_mlp_absent=true", "PASS"],
        ["Statistics", "Map is independent unit; paired tests and outer-map bootstrap", "PASS"],
        ["Abstract", f"{len(manuscript_abstract().split())} words", "PASS" if len(manuscript_abstract().split()) <= 250 else "FAIL"],
        ["Keywords", f"{len(KEYWORDS)}", "PASS"],
        ["Highlights", f"{len(HIGHLIGHTS)} bullets; max {max(map(len, HIGHLIGHTS))} characters", "PASS"],
        ["Anonymity", "Author identities omitted from manuscript; separate title page", "PASS WITH PLACEHOLDERS"],
        ["References", "DOI/official metadata register and claim map supplied", "PASS; access-limited full texts marked"],
        ["Generative images", "None used; plots redrawn from frozen Source Data", "PASS"],
        ["Data/code", "Anonymous/permanent repository links missing", "AUTHOR INPUT REQUIRED"],
        ["Funding/CRediT/COI", "No author-supplied facts", "AUTHOR INPUT REQUIRED"],
        ["Page limit", "To be verified after rendered layout", "PENDING RENDER AUDIT"],
        ["Current journal rule", "Recheck author guide and AI disclosure at submission", "AUTHOR CHECK"],
    ]
    add_table(doc, ["Audit area", "Evidence", "Status"], compliance_rows, [3.0, 10.0, 4.0])
    add_heading(doc, "6. Author Verification Queue", 1)
    add_bullets(doc, [
        "[AUTHOR CHECK 01] Author and affiliation identities.",
        "[AUTHOR CHECK 02] Hardware and software versions.",
        "[AUTHOR CHECK 03] Anonymous and permanent repository links/licences.",
        "[AUTHOR CHECK 04] Copernicus derivative redistribution terms.",
        "[AUTHOR CHECK 05] Funding, CRediT, competing interests, acknowledgements.",
        "[AUTHOR CHECK 06] Current EAAI AI disclosure wording at submission.",
    ])
    path = DOCS / "EAAI_traceability_and_compliance.docx"; doc.save(path); return path


def write_markdown_audits() -> None:
    style_lines = [
        "# Journal Style Profile",
        "",
        "- Journal: Engineering Applications of Artificial Intelligence",
        "- Article type: Original Research",
        "- Official format: single-column Word; double anonymised",
        "- Exemplar pool: 12 recent same-journal research articles",
        "- Use boundary: section functions and stable conventions only; no sentence-level imitation",
        "",
        "## Claim boundaries",
        "",
        "No real-flight validation, safety certification, out-of-range node generalisation, independent submask effects, or universal PPO dominance.",
    ]
    (EVIDENCE / "Journal_Style_Profile.md").write_text("\n".join(style_lines) + "\n", encoding="utf-8")
    matrix = ["# Evidence Matrix", "", "| Claim ID | Claim | Evidence | Type | Max strength | Boundary |", "|---|---|---|---|---|---|"]
    for row in CLAIM_MAP:
        matrix.append("| " + " | ".join(str(x).replace("|", "/") for x in row) + " |")
    (EVIDENCE / "Evidence_Matrix.md").write_text("\n".join(matrix) + "\n", encoding="utf-8")
    queue = ["# Author Verification Queue", "", *[f"- {x}" for x in [
        "[AUTHOR CHECK 01] Author and affiliation identities.",
        "[AUTHOR CHECK 02] Hardware and software versions.",
        "[AUTHOR CHECK 03] Anonymous and permanent repository links/licences.",
        "[AUTHOR CHECK 04] Copernicus derivative redistribution terms.",
        "[AUTHOR CHECK 05] Funding, CRediT, competing interests, acknowledgements.",
        "[AUTHOR CHECK 06] Current EAAI AI disclosure wording at submission.",
    ]]]
    (EVIDENCE / "Author_Verification_Queue.md").write_text("\n".join(queue) + "\n", encoding="utf-8")


def build_reproducibility_package() -> None:
    root = DELIVERY / "reproducibility"
    package = root / "anonymous_reproducibility_package"
    package.mkdir(parents=True, exist_ok=True)
    files_to_copy = [
        FROZEN / "formal_evaluation" / "results" / "final_audit_status.json",
        FROZEN / "formal_evaluation" / "evaluation_matrix_manifest.json",
        FROZEN / "analysis" / "pre_plot_statistics" / "descriptive_metrics.csv",
        FROZEN / "analysis" / "pre_plot_statistics" / "confirmatory_pairwise.csv",
        FROZEN / "analysis" / "pre_plot_statistics" / "confirmatory_omnibus.csv",
        FROZEN / "analysis" / "manuscript_training_aware_v2" / "training_dimension_scores.csv",
        FROZEN / "analysis" / "manuscript_preplot_closure_v5" / "hierarchical_bootstrap_summary.csv",
    ]
    for src in files_to_copy:
        shutil.copy2(src, package / src.name)
    scripts_dir = package / "manuscript_build_scripts"
    scripts_dir.mkdir(exist_ok=True)
    for src in (DELIVERY / "scripts").glob("*"):
        if src.is_file():
            shutil.copy2(src, scripts_dir / src.name)
    readme = (
        "# Anonymous reproducibility package\n\n"
        "This package is a manuscript-facing extract of the frozen v3.2.14 evidence. It does not rerun training or evaluation.\n\n"
        "## Integrity gates\n\n"
        "- final_results.jsonl remains in the frozen workspace and contains 21,648 records.\n"
        "- final_audit_status.json reports passed=true and ppo_mlp_absent=true.\n"
        "- Map is the independent statistical unit.\n"
        "- Post-hoc operational scores do not replace safe_weighted_coverage.\n\n"
        "## Copernicus reconstruction\n\n"
        "Use the official Copernicus DEM GLO-30 registry: https://registry.opendata.aws/copernicus-dem/. "
        "Raw assets are omitted unless redistribution is confirmed. Region identifiers and task manifests are retained in the frozen evaluation package.\n\n"
        "## Rebuilding manuscript artifacts\n\n"
        "Run the scripts in manuscript_build_scripts in numeric workflow order: evidence snapshot, literature register, workbooks, figures, documents, then QA. "
        "Training and evaluation commands are intentionally excluded from manuscript rebuilding.\n"
    )
    (package / "README.md").write_text(readme, encoding="utf-8")
    manifest = []
    for path in sorted(package.rglob("*")):
        if path.is_file():
            manifest.append({"path": str(path.relative_to(package)), "bytes": path.stat().st_size, "sha256": hashlib.sha256(path.read_bytes()).hexdigest()})
    (package / "MANIFEST.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> None:
    DOCS.mkdir(parents=True, exist_ok=True); EVIDENCE.mkdir(parents=True, exist_ok=True)
    outputs = [build_manuscript(), build_title_page(), build_supplement(), build_highlights(), build_cover_letter(), build_traceability()]
    write_markdown_audits()
    build_reproducibility_package()
    print(json.dumps({"documents": [str(x) for x in outputs], "abstract_words": len(manuscript_abstract().split()), "highlight_max_chars": max(map(len, HIGHLIGHTS))}, indent=2))


if __name__ == "__main__":
    main()
