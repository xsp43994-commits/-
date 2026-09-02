from __future__ import annotations

"""从v3中英文Markdown源稿生成EAAI单栏投稿包和英文双栏阅读预览。"""

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "documents"
FIGS = ROOT / "figures" / "english"
FIGS_ZH = ROOT / "figures" / "chinese"
OUT = ROOT / "deliverables"
QA = ROOT / "qa"

# 重要可调参数：EAAI投稿稿采用A4、单栏、TNR 11 pt、1.5倍行距和1 inch页边距。
BODY_PT = 11.0
LINE_MULTIPLE = 1.50
FIG_WIDTH = 6.20


MAIN_FIGURES = {
    "[FIGURE 1": [FIGS / "main" / "F01_method_and_evaluation_workflow.png"],
    "[FIGURE 2": [FIGS / "main" / "M01_priority_weighted_coverage_english.png", FIGS / "main" / "M03_priority_stratum_effects_english.png"],
    "[FIGURE 3": [FIGS / "main" / "M02_safety_and_return_effects_english.png", FIGS / "main" / "M04_resource_use_english.png"],
    "[FIGURE 4": [FIGS / "main" / "M05_online_planning_time_ECDF_repaired.png", FIGS / "supplementary" / "S02_quality_time_tradeoff_english.png"],
    "[FIGURE 5": [FIGS / "main" / "M06_training_curves_english.png", FIGS / "main" / "M07_training_stability_efficiency_english.png"],
    "[FIGURE 6": [FIGS / "main" / "M08_unseen_maps_and_DSM_transfer_english.png", FIGS / "showcase" / "V02_fixed_DSM_route_repaired.png"],
    "[FIGURE 7": [FIGS / "main" / "M09_two_layer_robustness_english.png"],
    "[FIGURE 8": [FIGS / "main" / "M10_ablation_effects_english.png"],
}

CAPTIONS = {
    "[FIGURE 1": "Figure 1. Return-aware PPO–Pointer planning and frozen evaluation workflow. The diagram describes the implemented sequence from mountain-road task representation to composite return-feasibility screening and the predefined evidence outputs; it is a non-generative schematic and does not add experimental evidence.",
    "[FIGURE 2": "Figure 2. Coverage and priority-stratum evidence. (a) Map-level safe priority-weighted coverage on 24 unseen synthetic maps and eight DSM maps; points are independent-map aggregates and vertical markers are medians. (b) Full-model minus ablation differences across priority strata; these stratified values are descriptive and do not replace the confirmatory map-level endpoint.",
    "[FIGURE 3": "Figure 3. Safe-return and resource outcomes. (a) PPO–Pointer minus comparator percentage-point effects for safe completion and depot return under known shifts and hidden model/perception mismatch; intervals are map-level 95% bootstrap intervals. (b) Median energy, range and mission-time utilization among safe routes; the dashed line denotes the budget limit.",
    "[FIGURE 4": "Figure 4. Online planning time and quality–time trade-off. (a) Empirical cumulative distributions of per-task online planning time over the frozen evaluation jobs. (b) Map-level D1 versus 95th-percentile planning time on synthetic and DSM tasks. Timing values are specific to the frozen software and hardware protocol and are not cross-platform latency guarantees.",
    "[FIGURE 5": "Figure 5. Corrected validation trajectories, training stability and sample efficiency. (a) Safe weighted coverage on the same fixed 108-task external validation set at 26 checkpoints over 3,000 episodes; thin curves are five seeds, heavy curves are medians and bands are interquartile ranges. (b) Corrected post-hoc D6 stability and D7 normalized validation-area-under-the-curve scores, together with direction-aligned component and budget-sensitivity values. These training dimensions supplement rather than redefine the confirmatory final-route endpoint.",
    "[FIGURE 6": "Figure 6. Cross-map performance and zero-shot DSM simulation transfer. (a) Map-level D1 estimates and 95% bootstrap intervals on unseen procedural maps and eight Copernicus DSM maps. (b) Terrain, road network, inspection points and representative routes for one fixed DSM task. The route panel is illustrative and non-inferential; DSM evaluation remains simulation, not flight validation.",
    "[FIGURE 7": "Figure 7. Two-layer robustness evaluation. D1 retention is reported separately for shifts revealed to the planner and hidden model/perception mismatches. The dashed reference denotes retention of 1.0; maps, not tasks or routes, are the inferential units.",
    "[FIGURE 8": "Figure 8. Four component ablations. Points are full-model minus ablation mean differences in map-level safe weighted coverage; horizontal intervals are 95% map-bootstrap intervals. Asterisked significance, where reported in the Source Data, follows Holm adjustment within the prespecified ablation family.",
}

MAIN_FIGURES_ZH = {
    "[FIGURE 1": [FIGS_ZH / "main" / "F01_workflow_zh.png"],
    "[FIGURE 2": [FIGS_ZH / "main" / "M01_zh.png", FIGS_ZH / "main" / "M03_zh.png"],
    "[FIGURE 3": [FIGS_ZH / "main" / "M02_zh.png", FIGS_ZH / "main" / "M04_zh.png"],
    "[FIGURE 4": [FIGS_ZH / "main" / "M05_online_planning_time_zh.png", FIGS_ZH / "supplementary" / "S02_zh.png"],
    "[FIGURE 5": [FIGS_ZH / "main" / "M06_zh.png", FIGS_ZH / "main" / "M07_zh.png"],
    "[FIGURE 6": [FIGS_ZH / "main" / "M08_zh.png", FIGS_ZH / "showcase" / "V02_DSM_route_zh.png"],
    "[FIGURE 7": [FIGS_ZH / "main" / "M09_zh.png"],
    "[FIGURE 8": [FIGS_ZH / "main" / "M10_zh.png"],
}

CAPTIONS_ZH = {
    "[FIGURE 1": "图1. 返航感知PPO–Pointer规划与冻结评价流程。该流程图描述从山区公路任务表示、复合返航可行性筛选到预定义证据输出的实现顺序；它是非生成式示意图，不增加实验性证据。",
    "[FIGURE 2": "图2. 覆盖率与优先级分层证据。（a）24张未见合成地图和8张DSM地图上的地图级安全优先级加权覆盖率；各点为独立地图聚合值，竖线为中位数。（b）各优先级层级中完整模型减消融模型的差值；这些分层值为描述性结果，不替代确认性地图级终点。",
    "[FIGURE 3": "图3. 安全返航与资源结果。（a）在已知偏移和隐藏模型/感知失配下，PPO–Pointer相对对比方法在安全完成和返回基地方面的百分点效应；区间为地图级95% bootstrap区间。（b）安全路线中的能量、航程和任务时间中位预算利用率；虚线表示预算上限。",
    "[FIGURE 4": "图4. 在线规划时间与质量–时间权衡。（a）冻结评价作业中逐任务在线规划时间的经验累积分布。（b）合成任务和DSM任务上的地图级D1与规划时间第95百分位。时间值仅适用于冻结软硬件协议，不是跨平台延迟保证。",
    "[FIGURE 5": "图5. 纠正后的验证轨迹、训练稳定性与样本效率。（a）3,000回合内26个检查点在同一固定108任务外部验证集上的安全加权覆盖率；细线表示5个种子，粗线表示中位数，阴影带表示四分位距。（b）纠正后的事后D6稳定性、D7归一化验证曲线下面积，以及统一优劣方向的组成和预算敏感性值。这些训练维度是确认性最终路线终点的补充，而不是重新定义该终点。",
    "[FIGURE 6": "图6. 跨地图性能与零样本DSM仿真迁移。（a）未见程序化地图和8张Copernicus DSM地图上的地图级D1估计及95% bootstrap区间。（b）一个固定DSM任务的地形、公路网络、巡检点和代表性路线。路线面板仅用于示意，不参与推断；DSM评价仍为仿真而非飞行验证。",
    "[FIGURE 7": "图7. 双层鲁棒性评价。对规划器可见的偏移和隐藏模型/感知失配分别报告D1保持率。虚线参考值为1.0；推断单位是地图，而不是任务或路线。",
    "[FIGURE 8": "图8. 4项组成部分消融。各点为地图级安全加权覆盖率中完整模型减消融模型的均值差；横向区间为95%地图bootstrap区间。若Source Data报告星号显著性，则其遵循预先设定消融族内的Holm校正。",
}

SUPP_FIGURES = [
    ("Figure S1", FIGS / "supplementary" / "S01_performance_profile_english.png", "Regret performance profile for all frozen algorithms. Each curve is an empirical cumulative distribution over the stated task set; the panel is descriptive and does not change the map-level inferential unit."),
    ("Figure S2", FIGS / "supplementary" / "S02_quality_time_tradeoff_english.png", "Coverage–time trade-off on synthetic and DSM tasks using D1 and the protocol-specific 95th-percentile planning time."),
    ("Figure S3", FIGS / "supplementary" / "S03_oracle_regret_cost_english.png", "Traditional-planner oracle-regret intervals versus planning time. Solver status and certification share must be read together with these values."),
    ("Figure S4", FIGS / "supplementary" / "S04_scenario_heatmap_english.png", "Scenario-stratified mean D1 for the three principal learning models. Scenario cells are descriptive strata, not independent replication units."),
    ("Figure S5", FIGS / "supplementary" / "S05_failure_modes_english.png", "Robustness and failure-mode rates under known shifts and hidden mismatch. Values are frozen task aggregates; safe and return rates are not flight-certification evidence."),
    ("Figure S6", FIGS / "supplementary" / "S06_seven_model_training_english.png", "Training-batch priority-weighted coverage for the seven paper-eligible learning variants. Curves summarize five seeds per model over 3,000 episodes; this quantity is neither external-validation coverage nor raw reward."),
    ("Figure S7", FIGS / "supplementary" / "S07_posthoc_score_english.png", "Seven normalized dimensions and the corrected post-hoc 100-point arithmetic composite. D6 and D7 use the formal 108-task validation traces; the composite remains exploratory, weight-dependent and is not the confirmatory endpoint."),
    ("Figure S8", FIGS / "supplementary" / "S08_weight_sensitivity_english.png", "Joint sensitivity of the PPO–Pointer first-place share to the combined D6+D7 weight and operational floor after correcting both training dimensions."),
    ("Figure S9", FIGS / "showcase" / "V01_fixed_synthetic_route_english.png", "Representative routes on one fixed synthetic task. The panel is illustrative and not used for statistical inference."),
    ("Figure S10", FIGS / "showcase" / "V02_fixed_DSM_route_repaired.png", "Representative routes on one fixed DSM task with elevation and road context. This is a zero-shot simulation-transfer illustration, not field validation."),
]

SUPP_FIGURES_ZH = [
    ("图S1", FIGS_ZH / "supplementary" / "S01_zh.png", "全部冻结算法的遗憾性能剖面。每条曲线是给定任务集合上的经验累积分布；该面板为描述性结果，不改变地图级推断单位。"),
    ("图S2", FIGS_ZH / "supplementary" / "S02_zh.png", "使用D1和协议特定规划时间第95百分位表示合成任务与DSM任务的覆盖率–时间权衡。"),
    ("图S3", FIGS_ZH / "supplementary" / "S03_zh.png", "传统规划器的Oracle遗憾区间与规划时间。必须结合求解器状态和认证比例解释这些数值。"),
    ("图S4", FIGS_ZH / "supplementary" / "S04_zh.png", "3种主要学习模型按情境分层的平均D1。情境单元格是描述性分层，不是独立重复单位。"),
    ("图S5", FIGS_ZH / "supplementary" / "S05_zh.png", "已知偏移和隐藏失配下的鲁棒性与失败模式比例。数值为冻结任务聚合；安全率和返航率不构成飞行认证证据。"),
    ("图S6", FIGS_ZH / "supplementary" / "S06_zh.png", "7种论文可用学习变体的训练批次优先级加权覆盖率。曲线汇总每个模型5个种子、3,000个回合的训练过程；该指标不是外部验证覆盖率，也不是原始reward。"),
    ("图S7", FIGS_ZH / "supplementary" / "S07_zh.png", "7个归一化维度及纠正后的事后100分算术综合得分。D6与D7使用正式108任务验证轨迹；该得分具有探索性并依赖权重，不是确认性终点。"),
    ("图S8", FIGS_ZH / "supplementary" / "S08_zh.png", "纠正两个训练维度后，PPO–Pointer第一名占比对D6+D7总权重和运行区间下限的联合敏感性。"),
    ("图S9", FIGS_ZH / "showcase" / "V01_zh.png", "一个固定合成任务的代表性路线。该面板仅用于示意，不用于统计推断。"),
    ("图S10", FIGS_ZH / "showcase" / "V02_DSM_route_zh.png", "带高程和公路情境的一个固定DSM任务代表性路线。该图属于零样本仿真迁移示意，而不是现场验证。"),
]


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])


def set_columns(section, count: int, space_twips: int = 360) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    node = cols[0] if cols else OxmlElement("w:cols")
    if not cols:
        sect_pr.append(node)
    node.set(qn("w:num"), str(count)); node.set(qn("w:space"), str(space_twips))


def set_run_fonts(run, chinese: bool = False, size: float | None = None) -> None:
    run.font.name = "Times New Roman"
    if size is not None:
        run.font.size = Pt(size)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体" if chinese else "Times New Roman")


def configure_document(doc: Document, running_header: str = "", chinese: bool = False, reading_proof: bool = False) -> None:
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.2677), Inches(11.6929)
    sec.top_margin = sec.bottom_margin = Inches(1.0 if not reading_proof else 0.58)
    sec.left_margin = sec.right_margin = Inches(1.0 if not reading_proof else 0.62)
    sec.header_distance = sec.footer_distance = Inches(0.35)
    set_columns(sec, 2 if reading_proof else 1, 320)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"; normal.font.size = Pt(8.6 if reading_proof else BODY_PT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体" if chinese else "Times New Roman")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE if reading_proof else WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.0 if reading_proof else LINE_MULTIPLE
    normal.paragraph_format.space_after = Pt(3 if reading_proof else 5)
    normal.paragraph_format.widow_control = True

    sizes = [("Title", 14 if reading_proof else 16), ("Heading 1", 10.4 if reading_proof else 14), ("Heading 2", 9.3 if reading_proof else 12), ("Heading 3", 8.8 if reading_proof else 11)]
    for name, size in sizes:
        st = styles[name]; st.font.name = "Times New Roman"; st.font.size = Pt(size); st.font.color.rgb = RGBColor(0, 0, 0)
        st.font.bold = True; st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei" if chinese else "Times New Roman")
        st.paragraph_format.keep_with_next = True; st.paragraph_format.space_before = Pt(6 if reading_proof else 9); st.paragraph_format.space_after = Pt(3 if reading_proof else 4)
        p_bdr = st._element.pPr.find(qn("w:pBdr")) if st._element.pPr is not None else None
        if p_bdr is not None:
            st._element.pPr.remove(p_bdr)

    if "Figure Caption" not in styles:
        fc = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        fc = styles["Figure Caption"]
    fc.font.name = "Times New Roman"; fc.font.size = Pt(7.4 if reading_proof else 9); fc.font.color.rgb = RGBColor(0, 0, 0)
    fc._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体" if chinese else "Times New Roman")
    fc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; fc.paragraph_format.space_after = Pt(8); fc.paragraph_format.keep_with_next = True

    header = sec.header.paragraphs[0]
    header.text = running_header if reading_proof else ""; header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if header.runs:
        header.runs[0].font.name = "Times New Roman"; header.runs[0].font.size = Pt(7.5); header.runs[0].font.color.rgb = RGBColor(70, 70, 70)
    add_page_field(sec.footer.paragraphs[0])


def add_inline(paragraph, text: str) -> None:
    """处理有限 Markdown 行内格式，避免把控制标记留在最终 DOCX。"""
    token = re.compile(r"(\*\*.*?\*\*|`.*?`|\*.*?\*)")
    pos = 0
    for match in token.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        value = match.group(0)
        if value.startswith("**"):
            r = paragraph.add_run(value[2:-2]); r.bold = True
        elif value.startswith("`"):
            r = paragraph.add_run(value[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(8.5)
        else:
            r = paragraph.add_run(value[1:-1]); r.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def sorted_reference_lines(lines: list[str]) -> list[str]:
    try:
        idx = lines.index("## References")
    except ValueError:
        return lines
    head = lines[:idx + 1]
    text = "\n".join(lines[idx + 1:]).strip()
    blocks = [b.strip() for b in re.split(r"\n\s*\n", text) if b.strip()]
    blocks.sort(key=lambda s: re.sub(r"[^A-Za-z]", "", s).lower())
    tail: list[str] = []
    for b in blocks:
        tail.extend(b.splitlines()); tail.append("")
    return head + [""] + tail


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    old = tbl_pr.find(qn("w:tblBorders"))
    if old is not None:
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge, val, size in [("top", "single", "8"), ("bottom", "single", "8"), ("insideH", "single", "3"), ("insideV", "nil", "0"), ("left", "nil", "0"), ("right", "nil", "0")]:
        node = OxmlElement(f"w:{edge}"); node.set(qn("w:val"), val); node.set(qn("w:sz"), size); node.set(qn("w:color"), "666666")
        borders.append(node)
    tbl_pr.append(borders)


def add_table(doc: Document, rows: list[list[str]], chinese: bool = False, reading_proof: bool = False) -> None:
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"; set_table_borders(table)
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.cell(i, j); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, row[j].strip() if j < len(row) else "")
            for run in p.runs:
                set_run_fonts(run, chinese, 7.0 if reading_proof else 8.5); run.bold = i == 0
        if i == 0:
            set_repeat_table_header(table.rows[i])
            for cell in table.rows[i].cells:
                tc_pr = cell._tc.get_or_add_tcPr(); tc_borders = OxmlElement("w:tcBorders")
                bottom = OxmlElement("w:bottom"); bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6"); bottom.set(qn("w:color"), "444444")
                tc_borders.append(bottom); tc_pr.append(tc_borders)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_picture(doc: Document, path: Path, width: float = FIG_WIDTH) -> None:
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))


def add_picture_pair(doc: Document, paths: list[Path], width: float = 3.0) -> None:
    table = doc.add_table(rows=1, cols=2); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr; old = tbl_pr.find(qn("w:tblBorders"))
    if old is not None: tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "insideH", "insideV", "left", "right"):
        node = OxmlElement(f"w:{edge}"); node.set(qn("w:val"), "nil"); borders.append(node)
    tbl_pr.append(borders)
    for idx, path in enumerate(paths):
        if not path.exists(): raise FileNotFoundError(path)
        p = table.cell(0, idx).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))


def add_main_figure(doc: Document, marker: str, chinese: bool = False, reading_proof: bool = False) -> None:
    figures = MAIN_FIGURES_ZH if chinese else MAIN_FIGURES
    captions = CAPTIONS_ZH if chinese else CAPTIONS
    key = next(k for k in figures if marker.startswith(k))
    if len(figures[key]) == 2 and not reading_proof:
        add_picture_pair(doc, figures[key], 3.00)
    else:
        for path in figures[key]:
            add_picture(doc, path, 6.1 if not reading_proof else 3.05)
    p = doc.add_paragraph(style="Figure Caption"); add_inline(p, captions[key])


def add_editable_equation(doc: Document, chinese: bool = False) -> None:
    """插入可在Word公式编辑器中直接修改的OMML公式及右侧编号。"""
    table = doc.add_table(rows=1, cols=2); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(5.4); table.columns[1].width = Inches(0.5)
    tbl_pr = table._tbl.tblPr; borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "insideH", "insideV", "left", "right"):
        node = OxmlElement(f"w:{edge}"); node.set(qn("w:val"), "nil"); borders.append(node)
    tbl_pr.append(borders)
    # 公式采用线性可编辑表示；不把数学表达式烘焙成图片。
    p = table.cell(0, 0).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    math_para = OxmlElement("m:oMathPara"); math_para.set(qn("m:jc"), "center")
    math = OxmlElement("m:oMath"); mr = OxmlElement("m:r"); mt = OxmlElement("m:t")
    mt.text = "SWC(π) = C_w(π), safe return with no hard-constraint violation; 0, otherwise"
    mr.append(mt); math.append(mr); math_para.append(math); p._p.append(math_para)
    q = table.cell(0, 1).paragraphs[0]; q.alignment = WD_ALIGN_PARAGRAPH.RIGHT; q.add_run("(1)")


def english_reference_lines() -> list[str]:
    lines = (DOCS / "EAAI_manuscript_source_v4.md").read_text(encoding="utf-8").splitlines()
    # 中文审阅稿必须继承英文投稿稿实际采用的同一排序，而不是源文件的历史顺序。
    lines = sorted_reference_lines(lines)
    idx = lines.index("## References")
    return lines[idx + 1:]


def markdown_to_doc(doc: Document, path: Path, main: bool = False, chinese: bool = False, reading_proof: bool = False) -> None:
    lines = path.read_text(encoding="utf-8").splitlines()
    if main and not chinese:
        lines = sorted_reference_lines(lines)
    if main and chinese and "[REFERENCES_FROM_ENGLISH]" in lines:
        marker = lines.index("[REFERENCES_FROM_ENGLISH]")
        lines = lines[:marker] + english_reference_lines() + lines[marker + 1:]
    i = 0
    first_title = True
    in_references = False
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1; continue
        if main and any(line.startswith(k) for k in MAIN_FIGURES):
            add_main_figure(doc, line, chinese, reading_proof); i += 1; continue
        if line.startswith("[EQUATION 1:"):
            add_editable_equation(doc, chinese); i += 1; continue
        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r"\s*:?-+:?\s*", c) for c in cells):
                    rows.append(cells)
                i += 1
            if rows: add_table(doc, rows, chinese, reading_proof)
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Title" if first_title else "Heading 1"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER if first_title else WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, line[2:]); first_title = False; i += 1; continue
        if line.startswith("## "):
            if line in ("## References", "## 参考文献"):
                doc.add_page_break()
                in_references = True
            p = doc.add_paragraph(style="Heading 1"); add_inline(p, line[3:]); i += 1; continue
        if line.startswith("### "):
            style = "Figure Caption" if re.match(r"### Table", line) else "Heading 2"
            p = doc.add_paragraph(style=style); add_inline(p, line[4:]); i += 1; continue
        if line.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3"); add_inline(p, line[5:]); i += 1; continue
        if re.match(r"^[-*] ", line):
            p = doc.add_paragraph(style="List Bullet"); add_inline(p, line[2:]); i += 1; continue
        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number"); add_inline(p, re.sub(r"^\d+\. ", "", line)); i += 1; continue

        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r"^(#{1,4} |[-*] |\d+\. |\|)", lines[i].strip()) and not (main and any(lines[i].startswith(k) for k in MAIN_FIGURES)) and not lines[i].startswith("[EQUATION 1:") and not re.match(r"^\*\*[^*]+:\*\*", lines[i].strip()):
            para_lines.append(lines[i].strip()); i += 1
        p = doc.add_paragraph(); add_inline(p, " ".join(para_lines))
        if in_references:
            p.paragraph_format.left_indent = Inches(0.25); p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE; p.paragraph_format.space_after = Pt(4)
            for run in p.runs: run.font.size = Pt(9.2 if not reading_proof else 7.2)
        if chinese:
            for run in p.runs: set_run_fonts(run, True)


def save_doc(md_name: str, out_name: str, header: str = "", main: bool = False, supplement: bool = False, chinese: bool = False, reading_proof: bool = False, compact: bool = False) -> Path:
    doc = Document(); configure_document(doc, header, chinese, reading_proof)
    if compact:
        normal = doc.styles["Normal"]; normal.font.size = Pt(10.2)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; normal.paragraph_format.line_spacing = 1.12
        normal.paragraph_format.space_after = Pt(3)
    markdown_to_doc(doc, DOCS / md_name, main=main, chinese=chinese, reading_proof=reading_proof)
    if supplement:
        doc.add_page_break(); doc.add_heading("补充图" if chinese else "Supplementary figures", level=1)
        for title, path, caption in (SUPP_FIGURES_ZH if chinese else SUPP_FIGURES):
            add_picture(doc, path, 6.1 if not reading_proof else 3.05)
            p = doc.add_paragraph(style="Figure Caption"); add_inline(p, f"{title}. {caption}")
    out = OUT / out_name; out.parent.mkdir(parents=True, exist_ok=True); doc.save(out)
    return out


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True); QA.mkdir(parents=True, exist_ok=True)
    produced = [
        save_doc("EAAI_manuscript_source_v4.md", "EAAI_manuscript_anonymized_v4.docx", main=True),
        save_doc("EAAI_manuscript_source_v4.md", "EAAI_manuscript_two_column_reading_proof_v4.docx", "Author-formatted reading proof — not publisher typeset", main=True, reading_proof=True),
        save_doc("EAAI_title_page_source_v4.md", "EAAI_title_page_v4.docx"),
        save_doc("EAAI_supplementary_source_v4.md", "EAAI_supplementary_material_v4.docx", supplement=True),
        save_doc("EAAI_highlights_source_v4.md", "EAAI_highlights_v4.docx"),
        save_doc("EAAI_cover_letter_source_v4.md", "EAAI_cover_letter_v4.docx", compact=True),
        save_doc("EAAI_manuscript_zh_v4.md", "EAAI_manuscript_zh_v4.docx", main=True, chinese=True),
        save_doc("EAAI_title_page_zh_v4.md", "EAAI_title_page_zh_v4.docx", chinese=True),
        save_doc("EAAI_supplementary_zh_v4.md", "EAAI_supplementary_material_zh_v4.docx", supplement=True, chinese=True),
        save_doc("EAAI_highlights_zh_v4.md", "EAAI_highlights_zh_v4.docx", chinese=True),
        save_doc("EAAI_cover_letter_zh_v4.md", "EAAI_cover_letter_zh_v4.docx", chinese=True),
    ]
    report = {
        "preset": "EAAI author manuscript v4",
        "named_override": "A4 single-column submission manuscript; separate English two-column reading proof",
        "body_pt": BODY_PT,
        "line_multiple": LINE_MULTIPLE,
        "documents": [str(p.relative_to(ROOT)) for p in produced],
    }
    (QA / "docx_build_profile_v4.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
